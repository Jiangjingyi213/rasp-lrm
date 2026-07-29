from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path("/Users/jiangjiang/Documents/Research/rasp-lrm/docs/RASP_LRM_Experiments_修订版_2026_07_29.docx")


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, size: float = 8.5, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    parts = str(text).split("\n")
    p = cell.paragraphs[0]
    p.alignment = align
    for i, part in enumerate(parts):
        if i:
            p.add_run().add_break()
        r = p.add_run(part)
        r.bold = bold
        r.font.size = Pt(size)
        r.font.name = "Times New Roman"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, top="000000", bottom="000000", inside="BFBFBF") -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "bottom"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "12")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), top if edge == "top" else bottom)
    for edge in ("insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), inside)
    for edge in ("left", "right"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "nil")


def set_cell_width(cell, width_cm: float) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_row_cant_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def add_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(9)
    r.font.name = "Hiragino Sans GB"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.55)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(text)
    r.font.size = Pt(10.5)
    r.font.name = "Hiragino Sans GB"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.size = Pt(8.5)
    r.italic = True
    r.font.color.rgb = RGBColor(80, 80, 80)
    r.font.name = "Hiragino Sans GB"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13 if level == 1 else 11)
    r.font.name = "Hiragino Sans GB"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")


def style_table(table, widths, header_rows=1) -> None:
    table.autofit = False
    set_table_borders(table)
    for ri, row in enumerate(table.rows):
        set_row_cant_split(row)
        if ri < header_rows:
            set_repeat_table_header(row)
        for ci, cell in enumerate(row.cells):
            if ci < len(widths):
                set_cell_width(cell, widths[ci])
            if ri < header_rows:
                set_cell_shading(cell, "EDEDED")


def add_main_table(doc: Document) -> None:
    headers = ["Ratio", "Method", "GSM8K", "MATH", "AMC", "GPQA", "ARC-C", "Avg.",
               "GSM8K", "MATH", "AMC", "GPQA", "ARC-C", "Avg."]
    rows = [
        ["20%", "Static", "73.84", "52.20", "30.00", "21.72", "79.10", "68.66", "pending", "pending", "pending", "pending", "pending", "pending"],
        ["", "Fixed-Global", "pending", "pending", "pending", "pending", "pending", "pending", "—", "—", "—", "—", "—", "—"],
        ["", "Wanda-C4", "pending", "pending", "pending", "pending", "pending", "pending", "—", "—", "—", "—", "—", "—"],
        ["", "RASP-LRM", "75.74", "58.00", "27.50", "27.78", "80.38", "71.14", "pending", "pending", "pending", "pending", "pending", "pending"],
        ["30%", "Static", "65.13", "42.60", "20.00", "19.19", "75.00", "61.85", "83.78", "pending", "37.50", "29.29", "88.05", "80.98†"],
        ["", "Fixed-Global", "pending", "pending", "pending", "pending", "pending", "pending", "—", "—", "—", "—", "—", "—"],
        ["", "Wanda-C4", "pending", "pending", "pending", "pending", "pending", "pending", "—", "—", "—", "—", "—", "—"],
        ["", "RASP-LRM", "70.20", "46.80", "20.00", "23.23", "76.45", "65.35", "83.62", "55.60", "55.00", "35.35", "89.33", "82.15†"],
    ]
    add_caption(doc, "表 1  主结果：不同剪枝预算下的准确率比较。数值为百分比，Avg. 为按样本数加权的平均准确率。")
    table = doc.add_table(rows=2, cols=len(headers))
    top = table.rows[0].cells
    top[0].merge(table.rows[1].cells[0])
    top[1].merge(table.rows[1].cells[1])
    top[2].merge(top[7])
    top[8].merge(top[13])
    set_cell_text(table.rows[0].cells[0], "Ratio", bold=True, size=7.4)
    set_cell_text(table.rows[0].cells[1], "Method", bold=True, size=7.4)
    set_cell_text(table.rows[0].cells[2], "Qwen3-1.7B", bold=True, size=7.6)
    set_cell_text(table.rows[0].cells[8], "Qwen3-4B", bold=True, size=7.6)
    for i, h in enumerate(headers[2:], start=2):
        set_cell_text(table.rows[1].cells[i], h, bold=True, size=7.1)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, bold=(val == "RASP-LRM"), size=6.8 if i >= 2 else 7.2)
        if row[1] == "RASP-LRM":
            for c in cells:
                set_cell_shading(c, "F5F5F5")
    style_table(table, [1.1, 2.05, 1.28, 1.28, 1.2, 1.2, 1.28, 1.2, 1.28, 1.28, 1.2, 1.2, 1.28, 1.2], header_rows=2)
    add_note(doc, "注：Static 表示整条推理轨迹共享同一套全局通道排序与固定剪枝率；RASP-LRM 表示本文方法。Fixed-Global 与 Wanda-C4 作为 Qwen3-1.7B 补充 baseline，目前保留占位，Qwen3-4B 不计划重复这两组补充 baseline，因此以 “—” 标记。† 表示 Qwen3-4B 的 static MATH-500 summary 尚未同步，因此 4B-T30 Avg. 暂按已完成 matched static 的四个任务加权；4B-T20 将在同步后补入。")


def add_stability_table(doc: Document) -> None:
    headers = ["Model", "Ratio", "Method", "Avg. Acc.", "Actual Prune", "Fallback", "Trunc."]
    rows = [
        ["Qwen3-1.7B", "20%", "Static", "68.66", "23.23", "12.67", "1.52"],
        ["", "20%", "RASP-LRM", "71.14", "22.51", "8.05", "0.50"],
        ["", "30%", "Static", "61.85", "34.06", "16.07", "3.13"],
        ["", "30%", "RASP-LRM", "65.35", "33.91", "16.07", "1.64"],
        ["Qwen3-4B", "30%", "Static", "80.98†", "36.36†", "4.84†", "1.14†"],
        ["", "30%", "RASP-LRM", "82.15†", "31.95†", "7.55†", "0.37†"],
    ]
    add_caption(doc, "表 2  主结果对应的剪枝率与稳定性指标。")
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8.2)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, bold=(val == "RASP-LRM"), size=8.0)
        if row[2] == "RASP-LRM":
            for c in cells:
                set_cell_shading(c, "F5F5F5")
    style_table(table, [2.0, 1.2, 2.1, 1.8, 1.8, 1.6, 1.5])
    add_note(doc, "注：Actual Prune 为实际触发的 MLP channel pruning 比例；Fallback 为进入 dense fallback 的比例；Trunc. 为输出截断率。")


def add_budget_table(doc: Document) -> None:
    headers = ["Setting", "Static Acc/Prune", "RASP-LRM Acc/Prune", "Δ Acc", "说明"]
    rows = [
        ["1.7B T20, 5-task", "68.66 / 23.23", "71.14 / 22.51", "+2.48", "full+priority 共同五任务"],
        ["1.7B T30, 5-task", "61.85 / 34.06", "65.35 / 33.91", "+3.50", "与主方法剪枝率接近"],
        ["Diagnostic T20", "67.79 / 22.71", "71.73 / 22.05", "+3.94", "同一 1040 样本严格 matched"],
        ["Diagnostic T30", "63.27 / 33.45", "67.69 / 33.54", "+4.42", "同一 1040 样本严格 matched"],
    ]
    add_caption(doc, "表 3  剪枝预算分析。T20/T30 diagnostic 用于检验预算变化下的趋势，不替代主结果表。")
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8.5)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=8.2, align=WD_ALIGN_PARAGRAPH.LEFT if i == 4 else WD_ALIGN_PARAGRAPH.CENTER)
    style_table(table, [3.4, 3.0, 3.2, 1.4, 6.0])
    add_note(doc, "注：Diagnostic 结果来自 GSM8K-200、MATH500-200、ARC-Easy-200、ARC-Challenge-200 和 BBH-selected-240；其作用是控制预算与样本集合，证明收益不只来自某一个剪枝率。")


def add_ablation_table(doc: Document) -> None:
    headers = ["Variant", "GSM8K", "MATH", "ARC-E", "ARC-C", "BBH", "Overall", "Prune", "FB", "Trunc."]
    rows = [
        ["Dense", "82.50\n0.00", "65.50\n0.00", "89.50\n0.00", "80.00\n0.00", "67.92\n0.00", "76.73", "0.00", "6.25", "0.19"],
        ["Static", "67.50\n36.15", "45.50\n34.15", "82.50\n33.74", "68.50\n31.27", "54.58\n34.85", "63.37", "34.06", "16.06", "3.46"],
        ["Fixed-Global", "65.00\n37.84", "41.00\n36.51", "83.50\n35.20", "72.00\n34.21", "51.67\n36.63", "62.21", "36.10", "17.12", "5.19"],
        ["Fixed-Stage", "65.00\n37.74", "40.50\n36.14", "85.00\n34.40", "70.00\n33.80", "49.58\n36.12", "61.54", "35.66", "18.46", "4.62"],
        ["Prior-Only", "69.50\n34.09", "43.50\n35.62", "80.50\n30.63", "71.00\n29.95", "56.25\n32.73", "63.85", "32.61", "84.23", "7.02"],
        ["No-Core", "67.00\n34.02", "51.00\n35.60", "84.50\n32.13", "73.50\n32.36", "52.50\n33.69", "65.19", "33.56", "18.65", "2.40"],
        ["RASP-LRM", "67.50\n33.90", "43.00\n35.80", "86.00\n31.99", "73.00\n32.30", "59.58\n33.23", "65.58", "33.44", "19.04", "1.83"],
    ]
    add_caption(doc, "表 4  五个数据集上的核心消融结果。每个数据集单元格为 Acc. / actual MLP pruning，Overall、Prune、FB 和 Trunc. 为五数据集加权结果。")
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, size=8.0)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, bold=(row[0] == "RASP-LRM"), size=7.6, align=WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER)
        if row[0] == "RASP-LRM":
            for c in cells:
                set_cell_shading(c, "F5F5F5")
    style_table(table, [2.4, 1.55, 1.55, 1.55, 1.55, 1.55, 1.4, 1.25, 1.1, 1.1])
    add_note(doc, "注：该表保留完整五数据集分项，避免只用 overall 掩盖任务差异。Fixed-Global/Fixed-Stage 检验固定 mask 的上限；Prior-Only 检验仅靠离线先验是否足够；No-Core 检验 protected core 的稳定作用。")


def main() -> None:
    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.top_margin = Cm(1.4)
    sec.bottom_margin = Cm(1.4)
    sec.left_margin = Cm(1.25)
    sec.right_margin = Cm(1.25)

    styles = doc.styles
    styles["Normal"].font.name = "Hiragino Sans GB"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")
    styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("4  Experiments（第一版中文草稿）")
    r.bold = True
    r.font.size = Pt(16)
    r.font.name = "Hiragino Sans GB"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Hiragino Sans GB")

    add_note(doc, "写作定位：本稿只写 Experiments 部分，因此不重复 Introduction 中已经出现的静态剪枝示意图与反事实 motivation 图；数据集、baseline 和 T20/T30 预算在 4.1 中简要交代。")

    add_heading(doc, "4.1 Datasets, Baselines, and Pruning Budgets", 1)
    add_body(doc, "我们主要在数学推理、科学问答和通用推理任务上评估剪枝后模型的可靠性。主结果采用五个共同任务：GSM8K、MATH-500、AMC2023、GPQA-Diamond 和 ARC-Challenge；完整 T30 设置还包含 AIME2024 与 AIME2025，用于检验更高难度竞赛数学场景。消融实验使用 GSM8K、MATH-500、ARC-Easy、ARC-Challenge 和 BBH-selected 构成的 diagnostic suite，以便在较小但任务类型多样的集合上分析各模块作用。所有结果报告每个数据集准确率，并给出按样本数加权的平均准确率。")
    add_body(doc, "主 baseline 为 Static，即 trajectory-global fixed mask。它使用与 RASP-LRM 相同的阶段化 prompt 和校准来源，但在整条推理轨迹中始终使用同一套全局通道排序与固定剪枝率。Wanda-C4 与 Fixed-Global 作为 Qwen3-1.7B 上的补充 baseline，目前在表格中保留位置；由于它们不作为 Qwen3-4B 的主要对照，4B 对应位置不展开展示。")
    add_body(doc, "T20 和 T30 表示两个剪枝预算区间，分别对应较低和较高的 MLP channel pruning 设置。由于 RASP-LRM 含有 protected core、阶段约束和 dense fallback，最终实际剪枝率不一定等于名义目标；因此正文统一报告 actual MLP channel pruning。当前实现施加的是 logical structured mask，尚未物理删除权重或重写推理内核，所以本文将其作为结构化剪枝质量与推理可靠性的验证，而不直接宣称端到端硬件加速。")

    add_heading(doc, "4.2 Main Results", 1)
    add_body(doc, "表 1 按剪枝率和方法组织结果，并将 Qwen3-1.7B 与 Qwen3-4B 作为横向模型分组展示。这样的排版使读者可以直接比较同一预算下不同模型规模的趋势，同时避免将尚未同步的 4B 补充 baseline 作为完整结论呈现。")
    add_main_table(doc)
    add_stability_table(doc)
    add_body(doc, "在 Qwen3-1.7B 上，RASP-LRM 在两个剪枝预算下均优于 matched Static。T20 设置中，RASP-LRM 在 22.51% actual MLP pruning 下达到 71.14% 平均准确率，相比 Static 的 68.66% 提升 2.48 个百分点；T30 设置中，RASP-LRM 在 33.91% pruning 下达到 65.35%，相比 Static 的 61.85% 提升 3.50 个百分点。若使用包含 AIME2024/AIME2025 的完整七数据集 T30 结果，RASP-LRM 为 64.15%，Static 为 60.75%，二者实际剪枝率分别为 33.97% 与 34.08%，提升为 3.41 个百分点。")
    add_body(doc, "Qwen3-4B 当前已同步 T30 的主要结果。已有 matched static 的四个任务上，RASP-LRM 的加权平均准确率为 82.15%，高于 Static 的 80.98%；同时截断率由 1.14% 降至 0.37%。但 MATH-500 的 Qwen3-4B Static summary 尚未同步，因此 4B 的完整五任务 aggregate 仍需补齐后再作为最终主结论。")

    add_heading(doc, "4.3 Effect of Pruning Budgets", 1)
    add_body(doc, "为了避免把收益误解为某个单一剪枝率下的偶然现象，我们进一步比较 T20 与 T30 两个预算区间。这里的 budget analysis 只回答一个受控问题：在相同或接近的 actual MLP pruning 下，stage-aware dynamic selection 是否仍然优于 static global mask。它不替代主结果表，也不用于重新选择方法超参数。")
    add_budget_table(doc)
    add_body(doc, "结果显示，在较低剪枝率和较高剪枝率下，RASP-LRM 都保持了相对于 Static 的正收益。尤其在 1040 题 diagnostic 设置中，T20 与 T30 的 static matched 对照分别提升 3.94 和 4.42 个百分点。这说明方法收益并不只是来自剪枝率偏低，而是来自阶段约束下的实例级通道选择。")

    add_heading(doc, "4.4 Ablation Study", 1)
    add_body(doc, "表 4 按五个 diagnostic 数据集展开核心消融结果。与只给 overall 的写法相比，分数据集结果能更清楚地展示不同组件在数学、科学推理和 BBH 子任务上的差异，同时仍保留整体准确率、剪枝率、fallback 与截断率。")
    add_ablation_table(doc)
    add_body(doc, "第一，固定化的剪枝结构不能充分替代在线选择。Fixed-Global 与 Fixed-Stage 的整体准确率分别为 62.21% 和 61.54%，均低于完整 RASP-LRM。这说明仅把校准统计固化为一套全局 mask，或进一步固化为若干阶段 mask，都不能稳定适配长链推理过程中的状态变化。该结果也回应了本文的 motivation：阶段敏感性确实存在，因此剪枝决策需要在阶段层面设定边界，但不能把阶段直接理解成一张固定子网络。")
    add_body(doc, "第二，离线先验需要当前实例信号修正。Prior-Only 的准确率为 63.85%，但 fallback 率达到 84.23%，说明仅依赖阶段校准得到的 prior 会频繁触发安全回退，难以单独承担在线剪枝决策。相比之下，RASP-LRM 在阶段先验之外引入 runtime activation，用当前样本最近窗口内的通道活动重新排序候选通道，从而把共享阶段知识转化为实例级选择。")
    add_body(doc, "第三，protected core 的作用主要是稳定推理路径。去除 protected core 后，准确率从 65.58% 降至 65.19%，截断率从 1.83% 升至 2.40%。这表明 protected core 并不是为了简单提高保留通道数量，而是在每个 reasoning stage 中保留一组高置信通道，降低动态剪枝误切关键推理路径的风险。总体来看，消融结果支持本文的核心设计：stage prior 提供阶段级安全边界，runtime activation 负责实例级细化，protected core 进一步保护长推理的连续性。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
