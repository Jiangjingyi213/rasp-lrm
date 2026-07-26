#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/jiangjiang/Documents/Research/rasp-lrm")
OUT = ROOT / "docs" / "RASP_LRM_消融实验论文正文_格式整理版.docx"


COLORS = {
    "blue": RGBColor(46, 116, 181),
    "dark_blue": RGBColor(31, 77, 120),
    "ink": RGBColor(34, 48, 60),
    "muted": RGBColor(101, 113, 125),
    "light_blue": "E8EEF5",
    "light_gray": "F2F4F7",
    "callout": "F4F8FB",
    "warning": "FFF7E6",
    "border": "D7DEE8",
}

CN_FONT = "Songti SC"


def set_run_font(run, ascii_font: str = CN_FONT, east_asia: str | None = None) -> None:
    east_asia = east_asia or ascii_font
    run.font.name = ascii_font
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), ascii_font)


def set_style_font(style, ascii_font: str = CN_FONT, east_asia: str | None = None) -> None:
    east_asia = east_asia or ascii_font
    style.font.name = ascii_font
    r_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), ascii_font)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: RGBColor | None = None, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    set_run_font(run)
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_cell_margins(cell, top=80, start=90, bottom=80, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DEE8", size="6"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    r._r.append(fld_char1)
    r._r.append(instr)
    r._r.append(fld_char2)
    r.font.size = Pt(9)
    r.font.color.rgb = COLORS["muted"]


def add_run(paragraph, text, bold=False, italic=False, color=None, size=None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    set_run_font(run)
    if color:
        run.font.color.rgb = color
    if size:
        run.font.size = Pt(size)
    return run


def add_para(doc, text="", style=None, before=None, after=None, line=1.18, align=None):
    p = doc.add_paragraph(style=style)
    if text:
        add_run(p, text)
    fmt = p.paragraph_format
    if before is not None:
        fmt.space_before = Pt(before)
    if after is not None:
        fmt.space_after = Pt(after)
    fmt.line_spacing = line
    if align:
        p.alignment = align
    return p


def add_callout(doc, title: str, body: str, fill="F4F8FB"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.3)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    set_table_borders(table, color="D7DEE8", size="6")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    add_run(p, title, bold=True, color=COLORS["dark_blue"], size=10.5)
    p2 = cell.add_paragraph()
    p2.paragraph_format.line_spacing = 1.15
    p2.paragraph_format.space_after = Pt(0)
    add_run(p2, body, size=10, color=COLORS["ink"])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    return p.add_run(text)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_run(p, text, size=10.5)


def add_code_block(doc, code: str):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F7F7F7")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    set_table_borders(table, color="E1E4E8", size="6")
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    for i, line in enumerate(code.splitlines()):
        if i:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, "Consolas", east_asia=CN_FONT)
        run.font.size = Pt(8.2)
        run.font.color.rgb = RGBColor(43, 58, 66)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def build_main_table(doc):
    headers = ["Variant", "Prior / mask", "Runtime", "Safety", "Acc. ↑", "Prune", "Fallback", "Trunc.", "Δ"]
    rows = [
        ["Dense", "--", "--", "--", "76.73", "0.00", "6.25", "0.19", "--"],
        ["Static", "global fixed", "✗", "✗", "63.37", "34.06", "16.06", "3.46", "0.00"],
        ["Fixed-Global", "global fixed", "✗", "stage ratio", "62.21", "36.10", "17.12", "5.19", "-1.16"],
        ["Fixed-Stage", "stage fixed", "✗", "stage ratio", "61.54", "35.66", "18.46", "4.62", "-1.83"],
        ["Prior-Only", "stage prior", "✗", "✓", "63.85", "32.61", "84.23", "7.02", "+0.48"],
        ["No-Core", "stage prior", "✓", "w/o core", "65.19", "33.56", "18.65", "2.40", "+1.82"],
        ["Uniform-Budget", "stage prior", "✓", "uniform", "65.87", "32.95", "20.00", "1.44", "+2.50"],
        ["RASP-LRM", "stage prior", "✓", "✓", "65.58", "33.44", "19.04", "1.83", "+2.21"],
        ["Dynamic-Global", "global prior", "✓", "✓", "66.35", "33.84", "4.62", "1.63", "+2.98"],
    ]
    p = add_para(doc, "表 X  核心消融结果。", before=4, after=2, line=1.0)
    p.runs[0].bold = True
    p.runs[0].font.color.rgb = COLORS["dark_blue"]
    cap = add_para(
        doc,
        "所有方法均在相同的 1040 个样本上评测，包含 GSM8K-200、MATH500-200、ARC-Easy-200、ARC-Challenge-200 和 BBH-selected-240。Actual pruning 表示实际执行的 MLP 通道剪枝比例；Δ 表示相对 Static 的准确率变化。",
        after=5,
        line=1.12,
    )
    cap.runs[0].font.size = Pt(9)
    cap.runs[0].font.color.rgb = COLORS["muted"]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    widths = [2.25, 2.25, 1.35, 1.55, 1.25, 1.25, 1.35, 1.15, 1.0]
    set_table_width(table, widths)
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_shading(cell, COLORS["light_blue"])
        set_cell_text(cell, h, bold=True, color=COLORS["dark_blue"], size=8.4)
        set_cell_margins(cell, top=90, bottom=90, start=70, end=70)
    for row in rows:
        cells = table.add_row().cells
        for j, item in enumerate(row):
            bold = row[0] == "RASP-LRM"
            set_cell_text(cells[j], item, bold=bold, color=COLORS["ink"], size=8.15)
            set_cell_margins(cells[j], top=75, bottom=75, start=65, end=65)
            if row[0] == "RASP-LRM":
                set_cell_shading(cells[j], "FFF2E8")
            elif row[0] in {"Prior-Only"}:
                set_cell_shading(cells[j], "FFF9E8")
            elif j == 0:
                set_cell_shading(cells[j], "F8FAFC")
    return table


def build_appendix_table(doc):
    headers = ["Dataset", "Static Acc.", "RASP Acc.", "Δ", "Static prune", "RASP prune", "主要观察"]
    rows = [
        ["ARC-Challenge", "68.50", "73.00", "+4.50", "31.27", "32.30", "动态选择明显优于 static"],
        ["ARC-Easy", "82.50", "86.00", "+3.50", "33.74", "31.99", "动态选择保留更多性能"],
        ["BBH-selected", "54.58", "59.58", "+5.00", "34.85", "33.23", "对复杂混合推理收益明显"],
        ["GSM8K", "67.50", "67.50", "0.00", "36.15", "33.90", "同准确率下剪枝略低"],
        ["MATH500", "45.50", "43.00", "-2.50", "34.15", "35.80", "diagnostic 子集上动态 mask 更敏感"],
    ]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_borders(table)
    widths = [2.35, 1.45, 1.35, 0.85, 1.35, 1.35, 4.8]
    set_table_width(table, widths)
    for j, h in enumerate(headers):
        set_cell_shading(table.rows[0].cells[j], COLORS["light_blue"])
        set_cell_text(table.rows[0].cells[j], h, bold=True, color=COLORS["dark_blue"], size=8.3)
    for row in rows:
        cells = table.add_row().cells
        for j, item in enumerate(row):
            set_cell_text(cells[j], item, size=8.1)
            if j == 6:
                cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_cell_margins(cells[j], top=80, bottom=80, start=75, end=75)
    doc.add_paragraph()


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    set_style_font(normal)
    normal.font.size = Pt(10.8)
    normal.font.color.rgb = COLORS["ink"]
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.22

    for name, size, color, before, after in [
        ("Heading 1", 16, COLORS["blue"], 18, 8),
        ("Heading 2", 13, COLORS["blue"], 14, 6),
        ("Heading 3", 11.5, COLORS["dark_blue"], 10, 4),
    ]:
        s = styles[name]
        set_style_font(s)
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    for lname in ["List Bullet", "List Number"]:
        s = styles[lname]
        set_style_font(s)
        s.font.size = Pt(10.3)
        s.paragraph_format.space_after = Pt(3)
        s.paragraph_format.line_spacing = 1.18


def main():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    configure_styles(doc)
    add_page_number(section)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    add_run(header, "RASP-LRM · Ablation Study Draft", color=COLORS["muted"], size=8.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    add_run(title, "RASP-LRM 消融实验论文正文草稿", bold=True, color=COLORS["ink"], size=20)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    add_run(subtitle, "用于 Experiments / Ablation Study 小节 · 中文整理版", color=COLORS["muted"], size=10.5)

    add_callout(
        doc,
        "写作定位",
        "主写 t30 core ablation：它与主结果约 34% actual MLP channel pruning 对齐。核心叙述不是“stage-specific prior 永远最优”，而是“固定 mask 不够，校准安全先验 + 当前样本 runtime activation 才是有效动态剪枝的来源”。",
    )

    add_heading(doc, "建议小节标题", 1)
    add_para(doc, "中文：消融实验：动态通道选择的来源", after=2)
    add_para(doc, "英文：Ablation Study: Where Does Dynamic Channel Selection Help?", after=8)

    add_heading(doc, "可直接放入论文的正文", 1)
    paragraphs = [
        "为了进一步分析 RASP-LRM 的性能来源，我们在五个推理数据集的 1040 个样本上进行了核心消融实验。所有剪枝方法均采用相同的显式阶段化推理协议，并在接近 33%--36% 的实际 MLP 通道剪枝率下进行比较。我们主要考察三个问题：第一，固定全局 mask 或固定阶段 mask 是否足以支撑长推理剪枝；第二，当前样本的 runtime activation 是否对通道选择有必要；第三，protected core 与阶段预算等安全控制是否主要影响剪枝稳定性。",
        "表 X 给出了总体结果。与使用单一 trajectory-global mask 的静态剪枝相比，RASP-LRM 在相近实际剪枝率下将总体准确率从 63.37% 提升到 65.58%，同时将截断率从 3.46% 降低到 1.83%。这一结果说明，推理过程中的 MLP 通道冗余并不能被一个全局固定排序充分刻画。更重要的是，固定的 stage-specific mask 并没有带来收益，其准确率仅为 61.54%，低于静态全局 mask。这表明，RASP-LRM 的优势并不是来自“离线得到几张阶段 mask 后机械切换”，而是来自推理时根据当前阶段和当前样本激活进行在线通道选择。",
        "runtime activation 是该动态选择中最关键的信号之一。在相同 stage-specific prior 条件下，RASP-LRM 相比固定阶段 mask 提升 4.04 个百分点；在 global prior 条件下，加入 runtime activation 的动态版本也比固定 global mask 高 4.14 个百分点。这说明仅使用校准集统计得到的通道重要性仍然不足，因为不同问题和不同生成上下文会改变当前真正活跃的 MLP 通道。与此一致，Prior-Only 虽然达到 63.85% 的准确率，但 fallback 率高达 84.23%，说明只依赖离线先验会频繁触发安全回退，不能作为一个可靠的在线剪枝策略。",
        "阶段信息的作用需要更谨慎地理解。消融结果并不支持“stage-specific prior 在所有任务上都是最强单独排序”这一过强结论；例如，global-prior dynamic variant 的总体准确率达到 66.35%。因此，我们将 reasoning stage 更准确地定义为一种风险控制信号：它决定当前阶段允许多激进地剪枝、多久刷新一次 mask、哪些通道应作为 protected core 被优先保留，以及在阶段协议异常时是否回退到 dense 计算。runtime activation 则负责在该阶段约束下，根据当前实例选择实际保留的通道。换言之，stage 回答“当前推理阶段怎样剪更安全”，runtime activation 回答“当前样本此刻哪些通道更重要”。",
        "安全组件的结果也支持这一解释。去掉 protected core 后，准确率从 65.58% 小幅下降到 65.19%，截断率从 1.83% 上升到 2.40%。这说明 protected core 不是单独贡献最大准确率提升的模块，而主要用于减少动态 mask 更新过程中误删阶段关键通道的风险。另一方面，Uniform-Budget 的准确率略高于当前手工阶段预算，说明现有阶段预算并非最终最优配置。我们因此将阶段预算作为 RASP-LRM 的安全控制接口，而不是把当前手工比例本身作为核心贡献。总体来看，消融结果支持本文的核心观点：在长链推理中，有效的结构化剪枝需要结合校准得到的安全先验与当前样本的在线激活证据，而不是依赖单一静态 mask 或固定阶段 mask 切换。",
    ]
    for txt in paragraphs:
        add_para(doc, txt, after=7, line=1.22)

    add_heading(doc, "主表：推荐放正文", 1)
    build_main_table(doc)

    add_heading(doc, "LaTeX 表格代码", 1)
    add_para(doc, "如果后续转 AAAI LaTeX，可直接使用以下 booktabs 版本；需要在导言区加入 booktabs 与 pifont。", after=4)
    code = r"""\usepackage{booktabs}
\usepackage{pifont}
\newcommand{\cmark}{\ding{51}}
\newcommand{\xmark}{\ding{55}}

\begin{table}[t]
\centering
\small
\setlength{\tabcolsep}{3.8pt}
\caption{Core ablation results on five reasoning subsets. All pruning variants are evaluated on the same 1040 examples. Actual pruning denotes the executed MLP channel pruning ratio. $\Delta$ is measured against the static global-mask baseline.}
\label{tab:core_ablation}
\begin{tabular}{lcccccc}
\toprule
Variant & Runtime & Safety & Acc. $\uparrow$ & Prune $\uparrow$ & Fallback & $\Delta$ \\
\midrule
Dense & -- & -- & 76.73 & 0.00 & 6.25 & -- \\
Static & \xmark & \xmark & 63.37 & 34.06 & 16.06 & 0.00 \\
Fixed-Global & \xmark & ratio & 62.21 & 36.10 & 17.12 & -1.16 \\
Fixed-Stage & \xmark & ratio & 61.54 & 35.66 & 18.46 & -1.83 \\
Prior-Only & \xmark & \cmark & 63.85 & 32.61 & 84.23 & +0.48 \\
No-Core & \cmark & w/o core & 65.19 & 33.56 & 18.65 & +1.82 \\
Uniform-Budget & \cmark & uniform & 65.87 & 32.95 & 20.00 & +2.50 \\
\textbf{RASP-LRM} & \cmark & \cmark & \textbf{65.58} & \textbf{33.44} & \textbf{19.04} & \textbf{+2.21} \\
Dynamic-Global & \cmark & \cmark & 66.35 & 33.84 & 4.62 & +2.98 \\
\bottomrule
\end{tabular}
\end{table}"""
    add_code_block(doc, code)

    add_heading(doc, "可选附录表：分数据集核心对比", 1)
    add_para(doc, "该表适合放附录或实验补充，不建议放正文主表，因为它会让读者过早纠结个别数据集波动。", after=5)
    build_appendix_table(doc)

    add_callout(
        doc,
        "正文可保留的边界说明",
        "各数据集上的收益并不完全一致：RASP-LRM 在 ARC 和 BBH 子集上提升明显，在 GSM8K 上与 static 持平，但在 MATH500 diagnostic 子集上低于 static。这说明长数学推理对动态 mask 更新更敏感，后续可针对高符号推理阶段采用更保守的 protected core 或刷新策略。",
        fill="FFF7E6",
    )

    add_heading(doc, "可选补充：低剪枝率 V2 ablation 的写法", 1)
    add_para(doc, "低剪枝率 ablation 不建议作为主消融，因为约 13%--15% 理论剪枝下各方法差距较小，不能最有力支撑主张。可以作为补充说明：", after=5)
    add_para(
        doc,
        "在较低剪枝强度下，动态剪枝与静态剪枝的差距较小。V2.4 dynamic ratio 在 13.20% 理论剪枝率下达到 74.33%，而 matched static global 在 14.18% 理论剪枝率下为 74.23%。不过，在相同 V2.4 stage ratios 下，runtime dynamic 仍分别比 fixed global 和 fixed stage-specific mask 高 2.02 和 2.98 个百分点。这说明在线激活选择在更激进的剪枝预算下更能体现价值。",
        after=7,
    )

    add_heading(doc, "不建议写进主文的说法", 1)
    warnings = [
        "不要写：“stage-specific prior 总是优于 global prior。”因为 Dynamic-Global 在 t30 core ablation 中达到 66.35%，高于当前 RASP-LRM 的 65.58%。",
        "不要写：“当前手工 stage budget 是最优的。”因为 Uniform-Budget 为 65.87%，略高于 RASP-LRM。",
        "不要把 Prior-Only 写成一个强 baseline。它 fallback 率为 84.23%，说明大量样本回退 dense，不能说明只靠 prior 就能稳定剪枝。",
        "不要声称已经证明真实 wall-clock speedup。当前表中报告的是 logical / actual MLP channel pruning 与准确率保持效果，除非另有硬件计时结果，否则不要写端到端加速。",
    ]
    for item in warnings:
        add_bullet(doc, item)

    add_heading(doc, "最稳的一句话结论", 1)
    add_callout(
        doc,
        "推荐结论",
        "消融实验表明，RASP-LRM 的核心收益来自校准先验与当前样本 runtime activation 的结合：离线阶段知识提供安全边界，在线激活证据决定实例级通道选择。固定全局 mask 或固定阶段 mask 都不足以刻画长推理中的动态通道重要性，而 protected core 与阶段预算主要承担风险控制作用。",
        fill="F4F8FB",
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
