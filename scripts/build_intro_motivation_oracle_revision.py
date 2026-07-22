from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from build_paper_story_docx import (
    BLUE,
    MID_GRAY,
    NAVY,
    add_caption,
    add_heading,
    add_para,
    set_run_font,
    style_document,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "RASP_LRM_Introduction_Motivation_Insight_含Oracle图正式版.docx"
FIG1 = (
    ROOT
    / "runs"
    / "01_motivation"
    / "motivation_analysis"
    / "paper_figures"
    / "fig1_reasoning_stage_sensitivity_heatmaps.png"
)
FIG2 = (
    ROOT
    / "runs"
    / "01_motivation"
    / "motivation_analysis"
    / "paper_figures"
    / "fig2_oracle_gap_diagnostic.png"
)


def reset_running_text(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    for run in list(header.runs):
        run._element.getparent().remove(run._element)
    run = header.add_run("RASP-LRM  |  Introduction, Motivation and Core Insight")
    set_run_font(run, size=8.5, bold=True, color=MID_GRAY)

    footer = section.footer.paragraphs[0]
    if footer.runs:
        footer.runs[0].text = "AAAI Introduction Draft  ·  "
        set_run_font(footer.runs[0], size=8.5, color=MID_GRAY)


def add_compact_numbered_item(doc, number, text):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Pt(18)
    paragraph.paragraph_format.first_line_indent = Pt(-18)
    paragraph.paragraph_format.space_after = Pt(3.5)
    paragraph.paragraph_format.line_spacing = 1.15
    marker = paragraph.add_run(f"{number}.  ")
    set_run_font(marker, size=10, bold=True, color=BLUE)
    body = paragraph.add_run(text)
    set_run_font(body, size=10)
    return paragraph


def build():
    for figure_path in (FIG1, FIG2):
        if not figure_path.exists():
            raise FileNotFoundError(figure_path)

    doc = Document()
    style_document(doc)
    reset_running_text(doc)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(3)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    run = title.add_run("RASP-LRM：面向长链推理的阶段约束动态结构化剪枝")
    set_run_font(run, size=17, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(10)
    run = subtitle.add_run("Introduction（Motivation、Insight 与方法概述）")
    set_run_font(run, size=10, bold=True, color=MID_GRAY)

    add_heading(doc, "1  Introduction", 1)

    # Paragraph 1 — research background and importance.
    add_para(
        doc,
        "大型推理模型（Large Reasoning Models, LRMs）通过生成较长的中间推导，显著扩展了语言模型处理数学、"
        "逻辑与多步决策任务的能力。然而，自回归解码需要在每个 token 上重复执行宽维 Transformer 计算，"
        "使长链推理的计算开销随生成长度持续累积。如何在不破坏推理可靠性的前提下降低逐 token 计算，"
        "因而成为大模型高效部署中的关键问题。结构化 MLP 通道剪枝能够直接减少参与计算的中间通道，"
        "并保留规则的张量形状与实现接口，具有重要的研究与应用价值。",
    )

    # Paragraph 2 — current paradigm and its implicit assumption.
    add_para(
        doc,
        "现有结构化剪枝主要沿两条路线发展：一类依据离线校准统计构造全局静态子网络，另一类利用 prompt、"
        "probe 或局部激活为当前输入选择子网络[1–6]。前者通过跨样本聚合获得稳定的重要性估计，后者则以输入"
        "相关信号提高适应性；二者均能在一定程度上避免无差别删除关键结构。其共同的隐含假设是，一次校准或"
        "一次输入级选择得到的通道集合，足以代表后续大部分生成轨迹，即通道重要性在一次推理内部近似稳定。",
    )

    # Paragraph 3 — fundamental limitation.
    add_para(
        doc,
        "尽管上述范式取得了良好效果，其决策粒度与长链推理的状态粒度并不完全匹配。模型在一次生成中会从问题"
        "理解转向推导、检验和最终作答，局部计算目标与通道依赖随上下文持续变化。因而，一张 trajectory-global "
        "mask 可能忽略轨迹内部的状态迁移；而简单地为每个阶段冻结一张 mask，又会把阶段标签视为通道需求的充分"
        "统计量，忽略同一阶段内由具体问题和局部上下文造成的实例差异。问题的本质不是剪枝策略是否“动态”，"
        "而是其动态决策能否跟随正在演化的推理状态。",
    )

    # Paragraph 4 — empirical discovery and motivation experiments.
    add_para(
        doc,
        "为检验这一问题，我们在稠密模型原本回答正确的推理轨迹上构造 104,280 个反事实结构化剪枝动作，并统计"
        "局部干预是否导致最终答案反转。总体 answer flip rate 为 45.8%，其中 MLP-channel 动作为 38.0%；图 1 "
        "进一步显示，反转风险随推理功能、剪枝模块和干预比例显著变化。更重要的是，图 2 基于同一反事实表后验计算的"
        "诊断性 oracle 显示：当动作选择由全局固定细化为每题选择、再细化为每个 problem-step 选择时，可识别出的"
        "最大反转率由 69.1% 提升至 77.4% 和 83.0%。这里的 oracle 选择事后最易导致反转的动作，仅用于度量风险"
        "异质性，而非可部署剪枝性能。一个进一步的反常识结果是，fixed stage-specific 的方向性消融并未优于 fixed "
        "global（71.35% 对 72.31%），而结合当前实例 recent activation 的 runtime dynamic 达到 74.33%。尽管这些"
        "变体的实际 recorded pruning 尚未严格匹配，结果共同表明：风险确实具有阶段结构，但阶段标签本身不足以"
        "唯一决定最优子网络。",
    )

    figure_paragraph = doc.add_paragraph()
    figure_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_paragraph.paragraph_format.space_before = Pt(3)
    figure_paragraph.paragraph_format.space_after = Pt(1)
    figure_paragraph.paragraph_format.keep_with_next = True
    shape = figure_paragraph.add_run().add_picture(str(FIG1), width=Inches(6.1))
    shape._inline.docPr.set("name", "Counterfactual pruning sensitivity across reasoning functions")
    shape._inline.docPr.set("title", "推理功能相关的反事实剪枝敏感性")
    shape._inline.docPr.set(
        "descr",
        "两幅热图展示自动划分的推理功能阶段在不同剪枝模块和剪枝比例下的答案反转率。",
    )
    add_caption(
        doc,
        "图 1 | 反事实剪枝风险随推理功能、剪枝模块与干预强度变化。图中阶段由自动规则用于细粒度事后诊断；"
        "在线方法使用可因果解析的显式阶段控制状态。该图证明风险的状态依赖性，不直接规定部署时的剪枝预算。",
    )

    oracle_paragraph = doc.add_paragraph()
    oracle_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    oracle_paragraph.paragraph_format.space_before = Pt(3)
    oracle_paragraph.paragraph_format.space_after = Pt(1)
    oracle_paragraph.paragraph_format.keep_with_next = True
    oracle_shape = oracle_paragraph.add_run().add_picture(str(FIG2), width=Inches(6.05))
    oracle_shape._inline.docPr.set("name", "Diagnostic static-prompt-step oracle gap")
    oracle_shape._inline.docPr.set("title", "更细粒度动作选择揭示步骤相关剪枝风险")
    oracle_shape._inline.docPr.set(
        "descr",
        "柱状图展示全局、每题和每个问题步骤的诊断性最大答案反转率，分别为百分之六十九点一、七十七点四和八十三点零。",
    )
    add_caption(
        doc,
        "图 2 | 更细粒度的动作选择揭示 problem-step 级剪枝风险结构。Static、prompt-level 与 step-level oracle "
        "分别在全数据集、每个问题和每个 problem-step 上后验选择反转率最高的动作。三者均由同一反事实表计算；"
        "数值衡量可识别的风险异质性，而非可部署剪枝策略的准确率或加速收益。",
    )

    # Paragraph 5 — key insight and research hypothesis.
    add_para(
        doc,
        "这些观察揭示，stage-aware 并不等价于 stage-fixed，输入自适应也不等价于推理状态自适应。我们的核心 "
        "insight 是：Stage selects the constraint; the instance selects the channels. 阶段信息更适合刻画当前推理"
        "状态下不应轻易破坏的稳定约束，而当前实例的近期激活则决定其余通道中哪些此刻可以删除。由此，剪枝决策"
        "应采用双时间尺度：离线阶段统计限定安全搜索空间，在线实例证据在该空间内完成实际选择。",
    )

    # Paragraph 6 — method overview, with each mechanism tied to the insight.
    add_para(
        doc,
        "基于这一 insight，我们提出 RASP-LRM，一种训练无关、轻量校准的动态结构化 MLP 通道剪枝方法。为了让"
        "剪枝决策跟随可观测的推理状态，模型首先通过显式 marker 形成可因果解析的阶段控制协议；为了将离线观察到"
        "的阶段风险转化为稳定约束，独立校准轨迹用于估计 stage-conditioned importance、阶段均值与 protected "
        "core，而不是冻结完整阶段 mask；为了保留同一阶段内的实例适应性，在线解码仅在 protected core 之外结合"
        "近期激活选择实际删除通道，并动态救回持续活跃的非核心通道。Dense warmup、低频 mask refresh、均值补偿"
        "与协议异常时的 dense fallback 进一步抑制短窗口波动。由此，每个模块都直接对应前述经验发现，而非独立"
        "堆叠的启发式技巧。",
    )

    # Paragraph 7 — results summary and contributions.
    add_para(
        doc,
        "在约 34% reported decode-only logical MLP channel pruning 下，RASP-LRM 相比近似剪枝率匹配的"
        "trajectory-global 静态基线，在 GSM8K 与 MATH-500 合计 1,819 道题上将准确率从 58.93% 提升至 63.77%"
        "（+4.84 个百分点），并在 7 个数据集 3,289 道题上从 60.75% 提升至 64.15%（+3.41 个百分点）。这些"
        "结果支持阶段条件保护与受约束实例动态选择的互补作用。本文的主要贡献概括如下：",
    )
    add_compact_numbered_item(
        doc,
        1,
        "通过大规模反事实诊断揭示长链推理剪枝风险的状态依赖性，并利用 static–prompt–step oracle gap 说明细粒度状态条件决策的必要性。",
    )
    add_compact_numbered_item(
        doc,
        2,
        "提出“阶段限定约束、实例选择通道”的双时间尺度观点，并据此构建具有 protected core 与在线动态救回机制的 RASP-LRM。",
    )
    add_compact_numbered_item(
        doc,
        3,
        "在近似匹配的 reported logical pruning 条件下验证相对静态基线的一致准确率优势，形成从现象、insight、机制到结果的证据闭环。",
    )

    note = add_para(
        doc,
        "说明：文中 [1–6] 对应 FLAP、GRIFFIN、Probe Pruning、GLASS、SEAP 与 OCP；oracle 数值来自同一反事实"
        "剪枝表并最大化 observed flip，因此只作为诊断性风险上界。reported pruning 指解码阶段逻辑 MLP 通道稀疏率，"
        "不等同于已测得的端到端加速。",
        first_indent=False,
        before=4,
        after=0,
        line=1.08,
    )
    for run in note.runs:
        set_run_font(run, size=8.5, color=MID_GRAY)

    props = doc.core_properties
    props.title = "RASP-LRM Introduction: Motivation, Insight, and Method Overview"
    props.subject = "AAAI-oriented Chinese Introduction revision with bounded oracle evidence"
    props.author = ""
    props.keywords = "RASP-LRM; introduction; oracle analysis; structured pruning; reasoning state"
    props.comments = "Seven-paragraph evidence-driven Introduction revision with a bounded diagnostic oracle figure."

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
