#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/jiangjiang/Documents/Research/rasp-lrm")
OUT = ROOT / "docs" / "RASP_LRM_消融实验_正式论文插入版.docx"

CN_FONT = "SimSun"
EN_FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(80, 80, 80)


def set_run_font(run, en_font: str = EN_FONT, cn_font: str = CN_FONT) -> None:
    run.font.name = en_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), en_font)
    r_fonts.set(qn("w:hAnsi"), en_font)
    r_fonts.set(qn("w:eastAsia"), cn_font)
    r_fonts.set(qn("w:cs"), en_font)


def set_style_font(style, en_font: str = EN_FONT, cn_font: str = CN_FONT) -> None:
    style.font.name = en_font
    r_pr = style._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), en_font)
    r_fonts.set(qn("w:hAnsi"), en_font)
    r_fonts.set(qn("w:eastAsia"), cn_font)
    r_fonts.set(qn("w:cs"), en_font)


def run(paragraph, text: str, bold=False, italic=False, size: float | None = None, color=BLACK):
    r = paragraph.add_run(text)
    set_run_font(r)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    return r


def para(doc, text: str = "", style=None, size: float = 10.5, after: float = 5, line: float = 1.15):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        run(p, text, size=size)
    return p


def set_cell_margins(cell, top=60, start=70, bottom=60, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(cell, text: str, bold=False, italic=False, size=8.2, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = run(p, text, bold=bold, italic=italic, size=size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    return r


def set_cell_border(cell, **kwargs):
    """
    Set cell borders. kwargs keys: top, bottom, left, right, insideH, insideV.
    Value is a dict, e.g. {"val": "single", "sz": "8", "color": "000000"}.
    """
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = "w:{}".format(edge)
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        if edge_data is None:
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), edge_data.get("val", "single"))
            element.set(qn("w:sz"), edge_data.get("sz", "6"))
            element.set(qn("w:space"), edge_data.get("space", "0"))
            element.set(qn("w:color"), edge_data.get("color", "000000"))


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    set_style_font(normal)
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, before, after in [
        ("Heading 1", 14, 10, 5),
        ("Heading 2", 12, 8, 4),
    ]:
        style = doc.styles[name]
        set_style_font(style)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = BLACK
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True


def add_formal_table(doc: Document) -> None:
    cap = para(doc, after=2, line=1.0)
    run(cap, "表 X  核心消融结果。", bold=True, size=9.2)
    run(
        cap,
        " 所有剪枝方法均在同一 1040 个样本上评测；Prune 表示实际 MLP 通道剪枝率，Δ 表示相对 Static 的准确率变化。",
        size=9.2,
    )

    headers = ["Variant", "Prior / mask", "Runtime", "Safety", "Acc. ↑", "Prune", "Fallback", "Trunc.", "Δ"]
    rows = [
        ["Dense", "--", "--", "--", "76.73", "0.00", "6.25", "0.19", "--"],
        ["Static", "global fixed", "No", "No", "63.37", "34.06", "16.06", "3.46", "0.00"],
        ["Fixed-Global", "global fixed", "No", "ratio", "62.21", "36.10", "17.12", "5.19", "-1.16"],
        ["Fixed-Stage", "stage fixed", "No", "ratio", "61.54", "35.66", "18.46", "4.62", "-1.83"],
        ["Prior-Only", "stage prior", "No", "Yes", "63.85", "32.61", "84.23", "7.02", "+0.48"],
        ["No-Core", "stage prior", "Yes", "w/o core", "65.19", "33.56", "18.65", "2.40", "+1.82"],
        ["Uniform-Budget", "stage prior", "Yes", "uniform", "65.87", "32.95", "20.00", "1.44", "+2.50"],
        ["RASP-LRM", "stage prior", "Yes", "Yes", "65.58", "33.44", "19.04", "1.83", "+2.21"],
        ["Dynamic-Global", "global prior", "Yes", "Yes", "66.35", "33.84", "4.62", "1.63", "+2.98"],
    ]

    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [2.35, 2.15, 1.35, 1.55, 1.1, 1.1, 1.25, 1.05, 0.9]
    set_table_width(table, widths)

    top_rule = {"val": "single", "sz": "12", "color": "000000"}
    mid_rule = {"val": "single", "sz": "8", "color": "000000"}
    bottom_rule = {"val": "single", "sz": "12", "color": "000000"}

    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]
        set_cell_text(cell, h, bold=True, size=7.7)
        set_cell_border(cell, top=top_rule, bottom=mid_rule)

    for i, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        is_last = i == len(rows)
        is_main = row[0] == "RASP-LRM"
        for j, item in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j in (0, 1, 3) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], item, bold=is_main, size=7.6, align=align)
            if is_last:
                set_cell_border(cells[j], bottom=bottom_rule)
            else:
                set_cell_border(cells[j])

    doc.add_paragraph()
    note = para(doc, after=4, line=1.05)
    run(note, "Note. ", italic=True, size=8.4, color=GRAY)
    run(
        note,
        "Prior-Only 的 fallback 率较高，说明仅依赖离线先验并不能形成稳定的在线剪枝策略；Dynamic-Global 的竞争性结果表明，reasoning stage 更适合作为风险控制信号，而不是单独的固定排序。",
        size=8.4,
        color=GRAY,
    )


def main() -> None:
    doc = Document()
    configure_doc(doc)

    title = doc.add_paragraph(style="Heading 1")
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(title, "消融实验", bold=True, size=14)

    p1 = para(doc)
    run(
        p1,
        "为了分析 RASP-LRM 的性能来源，我们在五个推理数据集的 1040 个样本上进行核心消融。所有剪枝方法均采用相同的显式阶段化推理协议，并在接近 33%--36% 的实际 MLP 通道剪枝率下比较。消融主要检验三点：固定 mask 是否足够、runtime activation 是否必要，以及 protected core 与阶段预算是否提供稳定性收益。",
        size=10.2,
    )

    p2 = para(doc)
    run(
        p2,
        "如表 X 所示，RASP-LRM 在相近剪枝率下将 Static 的准确率从 63.37% 提升到 65.58%，同时将截断率从 3.46% 降至 1.83%。相比之下，Fixed-Stage 只有 61.54%，说明收益并不来自简单切换几张离线固定的阶段 mask，而来自推理时结合阶段约束与当前样本激活的在线通道选择。Prior-Only 的 fallback 率达到 84.23%，进一步表明离线先验不能单独承担剪枝决策。",
        size=10.2,
    )

    p3 = para(doc)
    run(
        p3,
        "这些结果也提示我们应谨慎理解阶段信息的作用：stage-specific prior 并非在所有设置下都是最优排序，Dynamic-Global 也取得了竞争性结果。因此，本文将 reasoning stage 视为风险控制信号，用于控制剪枝预算、mask 刷新、protected core 和 fallback；runtime activation 则负责在当前实例中选择实际保留的 MLP 通道。总体而言，消融支持本文的核心观点：长链推理中的结构化剪枝需要校准先验与在线激活证据共同驱动，而不是依赖单一静态 mask。",
        size=10.2,
    )

    add_formal_table(doc)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
