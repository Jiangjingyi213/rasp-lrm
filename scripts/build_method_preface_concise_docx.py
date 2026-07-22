from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from build_paper_story_docx import (
    AMBER,
    AMBER_LINE,
    BLUE,
    MID_GRAY,
    NAVY,
    add_callout,
    add_caption,
    add_heading,
    add_para,
    set_run_font,
    style_document,
)


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "RASP_LRM_方法前置叙事_Motivation_Insight_AAAI精简版.docx"
FIG1 = (
    ROOT
    / "runs"
    / "01_motivation"
    / "motivation_analysis"
    / "paper_figures"
    / "fig1_reasoning_stage_sensitivity_heatmaps.png"
)


def reset_running_text(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    for run in list(header.runs):
        run._element.getparent().remove(run._element)
    r = header.add_run("RASP-LRM  |  Motivation and Core Insight")
    set_run_font(r, size=8.5, bold=True, color=MID_GRAY)

    footer = section.footer.paragraphs[0]
    if footer.runs:
        footer.runs[0].text = "AAAI 方法前置叙事  ·  "
        set_run_font(footer.runs[0], size=8.5, color=MID_GRAY)


def add_compact_bullet(doc, text, number=None):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.16
    prefix = f"{number}.  " if number is not None else "•  "
    r1 = p.add_run(prefix)
    set_run_font(r1, size=10, bold=True, color=BLUE)
    r2 = p.add_run(text)
    set_run_font(r2, size=10)
    return p


def build():
    if not FIG1.exists():
        raise FileNotFoundError(FIG1)

    doc = Document()
    style_document(doc)
    reset_running_text(doc)

    # Slightly tighter than the full story draft while remaining comfortable to read.
    normal = doc.styles["Normal"]
    normal.font.size = Pt(10.2)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.2

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(3)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("RASP-LRM：从推理阶段敏感性到受约束动态剪枝")
    set_run_font(r, size=17, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(9)
    r = subtitle.add_run("方法章节前置叙事：Motivation、Insight 与方法导出")
    set_run_font(r, size=10, bold=True, color=MID_GRAY)

    lead = doc.add_paragraph(style="Lead")
    lead.paragraph_format.first_line_indent = Pt(0)
    lead.paragraph_format.space_after = Pt(8)
    r = lead.add_run(
        "核心主张：长链推理中的通道需求既非全程固定，也不能由阶段标签唯一决定。"
        "阶段应限定安全搜索空间，而当前实例应决定其中的实际删除通道。"
    )
    set_run_font(r, size=10.8, bold=True, color=NAVY)

    add_heading(doc, "1  研究背景与现有缺口", 1)
    add_para(
        doc,
        "大型推理模型依靠较长的中间推导提升复杂问题求解能力，却也在自回归解码中反复执行稠密 MLP。"
        "结构化通道剪枝因此成为降低推理计算的重要路径。现有方法通常根据离线校准获得一张全局子网络，"
        "或利用 prompt、probe 与局部激活为一次生成选择子网络[1–6]。这些方法虽然具有不同程度的输入适应性，"
        "但大多默认一次选出的通道集合可代表后续大部分生成轨迹。",
    )
    add_para(
        doc,
        "该假设在长链推理中并不稳固。随着模型从问题理解转向推导、验证和最终作答，计算目标持续变化；"
        "一张 trajectory-global mask 可能忽略轨迹内部的状态迁移。一个自然修正是为每个阶段固定一张 mask，"
        "但这又把阶段标签误当成通道需求的充分统计量，忽略同一阶段内不同问题和局部上下文的差异。",
    )

    add_heading(doc, "2  反常识发现与 Motivation", 1)
    add_para(
        doc,
        "我们首先在稠密模型的正确推理轨迹上构造 104,280 个反事实剪枝动作，考察局部结构删除是否改变最终答案。"
        "总体 answer flip rate 为 45.8%，其中 MLP-channel 动作为 38.0%。图 1 进一步表明，答案反转率随推理功能、"
        "剪枝模块和干预比例明显变化。这一结果说明，剪枝风险并非沿生成过程均匀分布，因此全程复用同一通道排序缺乏充分依据。",
    )

    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(3)
    pic_p.paragraph_format.space_after = Pt(1)
    pic_p.paragraph_format.keep_with_next = True
    shape = pic_p.add_run().add_picture(str(FIG1), width=Inches(6.2))
    shape._inline.docPr.set("name", "Reasoning-stage counterfactual pruning sensitivity")
    shape._inline.docPr.set("title", "推理功能相关的反事实剪枝敏感性")
    shape._inline.docPr.set(
        "descr",
        "两幅热图展示五类事后功能阶段在不同剪枝模块与剪枝比例下的答案反转率。",
    )
    add_caption(
        doc,
        "图 1 | 反事实剪枝风险随推理功能、剪枝模块与干预强度变化。五阶段仅用于更细粒度的事后诊断；"
        "部署时采用可在生成过程中可靠解析的四阶段控制协议。该图证明风险异质性，不直接规定各阶段的剪枝预算。",
    )

    add_para(
        doc,
        "然而，风险具有阶段差异并不意味着“每阶段一张固定 mask”即可解决问题。"
        "在相同 nominal stage-ratio schedule 的方向性消融中，fixed stage-specific 的准确率为 71.35%，"
        "低于 fixed global 的 72.31%，且与 shuffled-stage 同为 71.35%；引入当前实例 recent activation 后，"
        "runtime dynamic 达到 74.33%。尽管这些变体的实际 recorded pruning 尚非严格一致，该结果仍揭示了一个"
        "反常识事实：stage-aware 不等于 stage-fixed，阶段身份本身不足以唯一决定最优子网络。",
    )
    add_callout(
        doc,
        "清晰的 Motivation",
        "我们需要的不是一组由阶段标签直接索引的固定 mask，而是一种双时间尺度决策："
        "用离线阶段统计提供稳定的安全边界，再用当前实例的近期激活在边界内完成动态选择。",
    )

    add_heading(doc, "3  核心 Insight 与方法导出", 1)
    add_callout(
        doc,
        "核心 Insight",
        "Stage selects the constraint; the instance selects the channels. "
        "阶段回答“哪些通道不能轻易改变”，当前实例回答“其余通道中哪些此刻可以删除”。",
    )
    add_para(
        doc,
        "基于这一 insight，我们提出 RASP-LRM，一种训练无关、轻量校准的动态结构化 MLP 通道剪枝方法。"
        "首先，模型通过显式 marker 生成可因果解析的四阶段控制状态；其次，独立校准轨迹用于估计 stage-conditioned "
        "importance、阶段均值和 protected core，而不是冻结完整阶段 mask；最后，在线解码依据当前阶段窗口内的"
        "近期激活，只在 protected core 之外选择实际删除通道。若某个非核心通道在当前实例中持续活跃，"
        "它会被动态救回。Dense warmup、低频 mask refresh、均值补偿和协议异常时的 dense fallback 共同抑制在线噪声。",
    )
    add_para(
        doc,
        "该设计逐一回应前述问题：stage-conditioned prior 捕获轨迹内部的稳定风险差异，protected core 防止高风险误删，"
        "recent activation 则保留同一阶段内的实例适应性。由此，RASP-LRM 避免全局静态策略的欠适应、阶段固定策略的"
        "过度简化，以及完全自由在线重选可能引入的短窗口波动。下一节将给出各模块的正式定义与实现。",
    )
    add_callout(
        doc,
        "【方法总览图待补】",
        "建议绘制一张紧凑双路径图：左侧为 Offline Calibration（stage-conditioned statistics 与 protected core），"
        "右侧为 Online Decoding（causal stage parser、recent activation、non-core selection 与安全回退），"
        "中央突出“阶段限定安全空间，实例决定实际 mask”。",
        fill=AMBER,
        border=AMBER_LINE,
    )

    add_heading(doc, "4  主要结果与贡献", 1)
    add_para(
        doc,
        "在约 34% reported decode-only logical MLP channel pruning 下，RASP-LRM 相比近似剪枝率匹配的"
        "trajectory-global 静态基线，在 GSM8K 与 MATH-500 合计 1,819 道题上将准确率从 58.93% 提升至 63.77%"
        "（+4.84 个百分点，+88 道题）；在全部 7 个数据集 3,289 道题上从 60.75% 提升至 64.15%"
        "（+3.41 个百分点，+112 道题）。在 MATH-500 上，动态方法即使具有更高的 reported pruning"
        "（36.21% 对 34.07%），准确率仍提升 4.20 个百分点。结果支持我们的核心判断："
        "稳定的阶段条件保护与受约束的实例动态选择缺一不可。这里的 pruning 指解码阶段逻辑 MLP 通道稀疏率，"
        "不等同于已测得的端到端加速。",
    )
    add_para(doc, "本文与方法部分相关的贡献可概括为三点：", first_indent=False, after=3)
    add_compact_bullet(
        doc,
        "通过大规模反事实诊断揭示推理剪枝风险的状态依赖性，并发现 stage-aware 不等价于 stage-fixed。",
        1,
    )
    add_compact_bullet(
        doc,
        "提出“阶段限定安全搜索空间、实例决定实际删除通道”的双时间尺度观点，并据此构建 RASP-LRM。",
        2,
    )
    add_compact_bullet(
        doc,
        "在近似匹配的 reported logical pruning 下验证了相对静态基线的一致准确率优势，闭合现象、机制与结果证据链。",
        3,
    )

    note = add_para(
        doc,
        "说明：文中 [1–6] 对应 FLAP、GRIFFIN、Probe Pruning、GLASS、SEAP 与 OCP；合并至正式论文时沿用统一参考文献编号。",
        first_indent=False,
        before=4,
        after=0,
        line=1.08,
    )
    for run in note.runs:
        set_run_font(run, size=8.5, color=MID_GRAY)

    props = doc.core_properties
    props.title = "RASP-LRM Method Preface: Motivation and Core Insight"
    props.subject = "Concise AAAI Chinese manuscript preface before Methodology"
    props.author = ""
    props.keywords = "RASP-LRM; motivation; insight; structured pruning; reasoning stages"
    props.comments = "Concise evidence-aligned pre-Method narrative."

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
