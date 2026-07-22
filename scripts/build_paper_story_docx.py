from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "RASP_LRM_论文故事_Motivation_Insight_正式中文稿.docx"
FIG1 = ROOT / "runs" / "01_motivation" / "motivation_analysis" / "paper_figures" / "fig1_reasoning_stage_sensitivity_heatmaps.png"

NAVY = "243746"
BLUE = "365E7D"
LIGHT_BLUE = "EAF1F6"
LIGHT_GRAY = "F4F6F8"
MID_GRAY = "66717C"
LINE = "B8C3CC"
AMBER = "FFF4D6"
AMBER_LINE = "D5A33E"
WHITE = "FFFFFF"
BLACK = "111111"


def set_run_font(run, ascii_name="Arial Unicode MS", east_asia="Arial Unicode MS", size=10.5,
                 bold=None, italic=None, color=BLACK):
    run.font.name = ascii_name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge not in kwargs:
            continue
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        spec = kwargs[edge]
        element.set(qn("w:val"), spec.get("val", "single"))
        element.set(qn("w:sz"), str(spec.get("sz", 4)))
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), spec.get("color", LINE))


def set_table_geometry(table, widths_dxa, indent_dxa=120):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths_dxa[idx] / 1440)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_run_font(run, size=9, color=MID_GRAY)


def add_para(doc, text="", style=None, align=None, first_indent=True,
             before=0, after=6, line=1.25, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.widow_control = True
    if first_indent and style is None:
        p.paragraph_format.first_line_indent = Pt(21)
    if align is not None:
        p.alignment = align
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_run_font(r2)
    else:
        r = p.add_run(text)
        set_run_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.keep_with_next = True
    r = p.add_run(text)
    set_run_font(r, ascii_name="Arial Unicode MS", east_asia="Arial Unicode MS",
                 size={1: 14, 2: 12, 3: 11}.get(level, 11), bold=True,
                 color=NAVY if level == 1 else BLUE)
    return p


def add_caption(doc, text, above=False):
    p = doc.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6 if above else 4)
    p.paragraph_format.space_after = Pt(4 if above else 8)
    p.paragraph_format.line_spacing = 1.1
    p.paragraph_format.keep_with_next = above
    r = p.add_run(text)
    set_run_font(r, size=9, color="38434D")
    return p


def add_callout(doc, title, body, fill=LIGHT_BLUE, border=BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360], indent_dxa=120)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, start=170, bottom=120, end=170)
    spec = {"val": "single", "sz": 8, "color": border}
    set_cell_border(cell, top=spec, start=spec, bottom=spec, end=spec)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(title)
    set_run_font(r, size=10, bold=True, color=NAVY)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    r2 = p2.add_run(body)
    set_run_font(r2, size=9.5, color="37434D")
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)
    return table


def style_document(doc):
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.78)
    sec.left_margin = Inches(0.9)
    sec.right_margin = Inches(0.9)
    sec.header_distance = Inches(0.32)
    sec.footer_distance = Inches(0.35)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial Unicode MS"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    heading_tokens = {
        1: (14, 14, 7, NAVY),
        2: (12, 10, 5, BLUE),
        3: (11, 7, 3, BLUE),
    }
    for level, (size, before, after, color) in heading_tokens.items():
        st = styles[f"Heading {level}"]
        st.font.name = "Arial Unicode MS"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    caption = styles["Caption"]
    caption.font.name = "Arial Unicode MS"
    caption._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    caption._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    caption.font.size = Pt(9)
    caption.font.italic = False

    if "Lead" not in styles:
        lead = styles.add_style("Lead", WD_STYLE_TYPE.PARAGRAPH)
    else:
        lead = styles["Lead"]
    lead.font.name = "Arial Unicode MS"
    lead._element.rPr.rFonts.set(qn("w:ascii"), "Arial Unicode MS")
    lead._element.rPr.rFonts.set(qn("w:hAnsi"), "Arial Unicode MS")
    lead._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial Unicode MS")
    lead.font.size = Pt(11)
    lead.font.color.rgb = RGBColor.from_string(NAVY)
    lead.paragraph_format.space_before = Pt(4)
    lead.paragraph_format.space_after = Pt(8)
    lead.paragraph_format.line_spacing = 1.28
    lead.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    header = sec.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(0)
    hr = hp.add_run("RASP-LRM  |  Motivation, Insight and Method Story")
    set_run_font(hr, size=8.5, bold=True, color=MID_GRAY)

    footer = sec.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_after = Pt(0)
    fr = fp.add_run("AAAI 中文论证稿  ·  ")
    set_run_font(fr, size=8.5, color=MID_GRAY)
    add_page_field(fp)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_result_table(doc):
    add_caption(doc, "表 1 | 约 34% reported decode-only logical MLP channel pruning 下的主要结果。", above=True)
    headers = ["评测范围", "方法", "准确率", "Reported pruning", "相对静态基线"]
    rows = [
        ["标准数学 full\nGSM8K + MATH-500 (n=1,819)", "Structured Dense", "76.86%", "0.00%", "-"],
        ["", "Static trajectory-global", "58.93%", "35.27%", "-"],
        ["", "RASP-LRM", "63.77%", "34.73%", "+4.84 pp / +88"],
        ["全部 7 个数据集\n(n=3,289)", "Structured Dense", "74.43%", "0.00%", "-"],
        ["", "Static trajectory-global", "60.75%", "34.08%", "-"],
        ["", "RASP-LRM", "64.15%", "33.97%", "+3.41 pp / +112"],
        ["MATH-500\n(n=500)", "Structured Dense", "61.20%", "0.00%", "-"],
        ["", "Static trajectory-global", "42.60%", "34.07%", "-"],
        ["", "RASP-LRM", "46.80%", "36.21%", "+4.20 pp / +21"],
    ]
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 2300, 1450, 1750, 1660])
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        c = hdr.cells[i]
        c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(c, NAVY)
        set_cell_margins(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=8.6, bold=True, color=WHITE)
    for ridx, values in enumerate(rows):
        row = table.add_row()
        for cidx, value in enumerate(values):
            c = row.cells[cidx]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(c, top=78, bottom=78)
            if ridx in (0, 3, 6):
                set_cell_shading(c, LIGHT_GRAY)
            if values[1] == "RASP-LRM":
                set_cell_shading(c, LIGHT_BLUE)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if cidx in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(value)
            set_run_font(r, size=8.5, bold=(values[1] == "RASP-LRM"), color=BLACK)
    note = add_para(doc,
        "注：pruning 为按样本汇总的解码阶段逻辑 MLP 通道稀疏率，不包含 dense prefill，也不等同于已测得的端到端加速。Static 与 dynamic 的汇总计数口径并非完全同义，因此正文使用 approximately matched reported logical pruning。",
        first_indent=False, after=7, line=1.1)
    for run in note.runs:
        set_run_font(run, size=8.5, color=MID_GRAY)


def add_ablation_table(doc):
    caption = add_caption(doc, "表 2 | 固定阶段身份不足以替代实例阶段动态选择。", above=True)
    caption.paragraph_format.page_break_before = True
    headers = ["策略", "准确率", "Recorded pruning", "关键含义"]
    rows = [
        ["Runtime dynamic", "74.33%", "13.20%", "Stage prior + current-instance activation"],
        ["Fixed global", "72.31%", "14.98%", "整条轨迹复用全局 mask"],
        ["Fixed stage-specific", "71.35%", "14.71%", "每阶段冻结一张 mask"],
        ["Fixed shuffled-stage", "71.35%", "15.42%", "打乱阶段与 mask 的对应关系"],
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    set_table_geometry(table, [2200, 1350, 1700, 4110])
    set_repeat_table_header(table.rows[0])
    for i, text in enumerate(headers):
        c = table.rows[0].cells[i]
        set_cell_shading(c, NAVY)
        set_cell_margins(c)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, size=8.6, bold=True, color=WHITE)
    for values in rows:
        row = table.add_row()
        for cidx, value in enumerate(values):
            c = row.cells[cidx]
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(c, top=85, bottom=85)
            if values[0] == "Runtime dynamic":
                set_cell_shading(c, LIGHT_BLUE)
            p = c.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if cidx in (0, 3) else WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.05
            r = p.add_run(value)
            set_run_font(r, size=8.5, bold=(values[0] == "Runtime dynamic"))
    note = add_para(doc,
        "注：该消融采用相同的 V2.4 nominal stage-ratio schedule，但实际 recorded pruning 并非严格一致，因此它用于支持“固定 stage mask 不足”的方向性结论，而非精确的同稀疏率因果估计。",
        first_indent=False, after=7, line=1.1)
    for run in note.runs:
        set_run_font(run, size=8.5, color=MID_GRAY)


def build():
    if not FIG1.exists():
        raise FileNotFoundError(FIG1)
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(4)
    title.paragraph_format.space_after = Pt(5)
    title.paragraph_format.keep_with_next = True
    r = title.add_run("RASP-LRM：面向大型推理模型的显式阶段校准与受约束动态结构化剪枝")
    set_run_font(r, size=18, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(11)
    r = subtitle.add_run("Motivation、核心 Insight 与论文故事正式中文稿")
    set_run_font(r, size=10.5, bold=True, color=MID_GRAY)

    add_heading(doc, "摘要", 1)
    add_para(doc,
        "长链推理拓展了大型语言模型解决复杂问题的能力，但也使稠密 MLP 在解码过程中被反复调用。现有结构化剪枝通常在离线校准后冻结全局子网络，或根据 prompt / probe 为一次生成选择子网络；二者都隐含假设剪枝决策在后续轨迹中近似稳定。我们通过 104,280 个反事实剪枝动作观察到，答案对剪枝的敏感性随推理功能、剪枝模块与干预强度呈现明显差异。更反常识的是，直接为不同阶段冻结不同 mask 并不能解决该问题：在相同 nominal stage-ratio schedule 下，fixed stage-specific mask 甚至低于 fixed global mask，而加入当前实例激活后性能明显恢复。由此我们提出核心观点：reasoning stage 应约束安全搜索空间，而不应直接决定最终子网络。基于这一观点，RASP-LRM 将模型自生成的显式阶段标记转化为可因果解析的控制状态，利用离线阶段条件统计建立 protected core，并在非核心空间内依据当前阶段的近期激活动态选择实际删除通道。实验表明，在约 34% reported decode-only logical MLP pruning 下，RASP-LRM 相比近似剪枝率匹配的静态基线在 GSM8K 与 MATH-500 合计上提升 4.84 个百分点，并在七个数据集 3,289 道题上提升 3.41 个百分点。结果说明，阶段条件稳定性与受约束实例适应性相结合，比全局静态或阶段固定的剪枝策略更适合长链推理。")
    kw = add_para(doc, "关键词：大型推理模型；动态结构化剪枝；推理阶段；反事实分析；MLP 通道剪枝",
                  first_indent=False, after=9)
    for run in kw.runs:
        set_run_font(run, size=9.5, bold=True, color=MID_GRAY)

    add_heading(doc, "1  引言", 1)
    add_para(doc,
        "大型推理模型通过生成中间推导、检查与修正过程，在数学推理和复杂问答任务上表现出更强的问题求解能力。然而，推理轨迹的增长也意味着模型需要在自回归解码中反复执行 Transformer MLP；在长答案场景下，这部分稠密计算会持续累积。结构化通道剪枝能够以规则的通道粒度减少参与计算的中间维度，因此成为模型压缩与高效推理的重要方向。")
    add_para(doc,
        "现有训练无关结构化剪枝主要沿两条路线发展。静态方法根据离线校准数据估计通道重要性，并在所有输入和生成位置复用同一子网络，例如 FLAP 通过激活波动与输出权重构造结构化重要性并使用均值补偿[1]。随后，GRIFFIN、Probe Pruning、GLASS、SEAP 与 OCP 等方法进一步引入 prompt、probe、任务或当前 batch 信息，使剪枝决策能够适应输入差异[2-6]。尽管适应粒度不同，这些方法通常仍在一次决策后复用所选子网络，因而默认当前输入的重要通道能够代表后续大部分生成过程。")
    add_para(doc,
        "这一假设在长链推理中并不稳固。问题理解、计划形成、推导、验证和最终作答具有不同的语义目标；即使输入保持不变，模型在生成轨迹中的计算需求仍可能迁移。我们的反事实诊断进一步表明，剪枝动作是否破坏最终答案取决于其发生的推理功能、模块和强度。更重要的是，阶段差异并不意味着可以简单地为每个阶段冻结一张 mask：在现有消融中，fixed stage-specific mask 未能超过 fixed global mask，而结合当前实例近期激活的动态策略获得更高准确率。这个反常识结果表明，重要通道既不是全局固定的，也不能由 stage 标签唯一确定。")

    lead = doc.add_paragraph(style="Lead")
    lead.paragraph_format.first_line_indent = Pt(0)
    rr = lead.add_run("我们的核心 insight 是：阶段应当约束剪枝策略，而不应直接决定被剪子网络。离线阶段统计负责回答“哪些通道不能轻易改变”，当前实例证据负责回答“剩余空间中哪些通道此刻可以改变”。")
    set_run_font(rr, size=11, bold=True, color=NAVY)

    add_para(doc,
        "据此，我们提出 RASP-LRM（Reasoning-Stage Adaptive Structured Pruning for Large Reasoning Models）。该方法首先要求模型显式生成有序阶段标记，并将其解析为可观察、可验证且具有 dense fallback 的因果控制状态；随后在独立校准轨迹上估计 stage-conditioned WIFV-inspired prior、阶段均值与 protected core；推理时，recent activation 只在保护核心之外更新通道可剪性，从而在稳定先验与实例适应之间形成受约束的动态选择。Warmup、低频 refresh、mask cache、mean compensation 与协议回退进一步限制在线估计噪声。")
    add_para(doc, "本文的主要贡献如下：", first_indent=False, after=3)
    contribs = [
        "我们通过反事实剪枝分析揭示了长链推理中的剪枝风险异质性，并进一步发现 stage-aware 并不等价于 stage-fixed：阶段固定子网络不足以替代当前实例证据。",
        "我们将 reasoning stage 从固定 mask 的索引重新定义为可验证的安全控制条件，提出“stage-conditioned protected prior + constrained instance adaptation”的动态结构化剪枝框架。",
        "在约 34% reported decode-only logical MLP pruning 下，RASP-LRM 相比近似匹配的 trajectory-global 静态基线，在核心数学集合上提升 4.84 个百分点，并在七个数据集上提升 3.41 个百分点。",
    ]
    for item in contribs:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.18
        r = p.add_run(item)
        set_run_font(r, size=10.2)

    section2 = add_heading(doc, "2  经验动机：推理剪枝不是全程平稳过程", 1)
    section2.paragraph_format.page_break_before = True
    add_heading(doc, "2.1  功能阶段相关的反事实敏感性", 2)
    add_para(doc,
        "我们首先不预设具体部署策略，而是从稠密模型的正确推理轨迹出发，构造反事实剪枝动作，并观察在保持题目与前缀不变时，局部结构化干预是否改变最终答案。该诊断覆盖 GSM8K 与 MATH-500，共包含 104,280 个反事实动作；总体 answer flip rate 为 45.8%，其中 MLP-channel 动作为 38.0%。这说明相当一部分看似局部的结构删除能够沿后续生成传播并改变最终结论。")
    add_para(doc,
        "为分析风险随推理功能的变化，我们将已有轨迹自动划分为 understanding、planning、derivation、verification 与 final 五类功能片段。该五阶段划分仅用于事后诊断，目的是在不预先绑定方法协议的情况下获得更细的分析分辨率；实际部署控制采用更紧凑、可在生成时因果解析的四阶段 marker 协议。两者分别回答“风险是否随推理功能变化”和“如何获得可靠的在线控制状态”，因此不要求逐项一一对应。这种分离也避免了用方法自身定义的阶段协议反向制造 motivation。")

    pic_p = doc.add_paragraph()
    pic_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pic_p.paragraph_format.space_before = Pt(5)
    pic_p.paragraph_format.space_after = Pt(2)
    pic_p.paragraph_format.keep_with_next = True
    shape = pic_p.add_run().add_picture(str(FIG1), width=Inches(6.45))
    shape._inline.docPr.set("name", "Reasoning-stage pruning sensitivity heatmaps")
    shape._inline.docPr.set("title", "功能阶段相关的反事实剪枝敏感性")
    shape._inline.docPr.set(
        "descr",
        "两幅热图分别展示五类推理功能阶段在不同剪枝模块和不同剪枝比例下的答案反转率。",
    )
    add_caption(doc,
        "图 1 | 推理剪枝敏感性随功能阶段、剪枝模块与干预强度变化。左图给出五类事后功能阶段与结构化剪枝单元的 answer flip rate；右图给出不同阶段与剪枝比例的对应结果。该图用于证明风险非均匀性，不将各阶段绝对高低直接解释为部署预算，也不假设五类诊断标签与四阶段控制协议一一对应。")

    add_para(doc,
        "图 1 显示，answer flip 并不由剪枝比例单独决定，而是同时受推理功能和剪枝单元影响。例如 verification 片段在多个模块和比例下表现出较高敏感性，而不同模块的干预风险也存在明显差异。由于不同片段的干预位置和剩余生成长度并不完全相同，我们不把某个阶段的绝对 flip rate 解释为固定安全预算；这里可靠的结论是，推理轨迹中的剪枝风险具有状态依赖性，因而一张 trajectory-global mask 难以覆盖全部生成状态。")

    add_heading(doc, "2.2  反常识现象：Stage-aware 不等于 Stage-fixed", 2)
    add_para(doc,
        "一个直接但过于简单的响应是：为每个推理阶段各自校准一张固定 mask，并在阶段切换时更换子网络。如果 stage 足以完全决定通道需求，这一策略应当优于全局固定 mask。然而，现有 V2.4 消融给出了相反结果。在相同 nominal stage-ratio schedule 下，fixed stage-specific mask 的准确率为 71.35%，低于 fixed global 的 72.31%，并与 shuffled-stage 对照持平；加入当前实例 recent activation 后，runtime dynamic 策略达到 74.33%。虽然这些变体的 recorded pruning 尚非严格一致，但结果清楚表明：仅仅知道当前 stage 并冻结一张阶段 mask，不足以解释完整动态方法的收益。")
    add_ablation_table(doc)
    add_para(doc,
        "这一结果改变了问题的定义。真正需要建模的不是“当前属于哪个阶段，因此应该切换到哪张 mask”，而是“当前阶段提供了怎样的稳定风险结构，以及当前实例正在该结构中调用哪些通道”。阶段之间存在可测差异，但它们同时共享较大的重要通道核心；在共享核心之外，不同问题和局部推理状态仍会改变实际活跃通道。")

    add_heading(doc, "2.3  核心 Insight：阶段定义安全搜索空间，实例决定实际子网络", 2)
    add_callout(doc,
        "核心 Insight",
        "长链推理需要两种互补信息：稳定的 stage-conditioned protection skeleton 与当前实例的 recent activation。前者限制在线决策不能触碰的高风险通道，后者只在低风险剩余空间内决定实际删除集合。换言之，Stage selects the constraint; the instance selects the channels。")
    add_para(doc,
        "该 insight 同时解释了四种策略的局限。全局静态 mask 忽略生成过程内部的状态迁移；prompt-level mask 将初始输入当作整条轨迹的充分统计量；stage-fixed mask 忽略同一阶段内的实例差异；完全自由的 token-wise 重选则容易被短窗口噪声驱动。更合理的设计不是在静态与动态之间二选一，而是把稳定校准先验变成动态选择的约束条件。")

    add_callout(doc,
        "【图 2 待补：两时间尺度非平稳性的机制证据】",
        "建议投稿前补充一个三面板定量图：(a) stage-pair prior Jaccard / Spearman，显示共享核心与阶段修正；(b) 不同实例在同一 stage 内的 runtime-mask Jaccard 或 churn；(c) runtime rescue rate，即近期激活从离线待删集合中救回的通道比例。该图将直接闭合“阶段提供骨架、实例决定修正”的证据链。",
        fill=AMBER, border=AMBER_LINE)

    add_heading(doc, "3  RASP-LRM：受阶段约束的实例动态结构化剪枝", 1)
    add_heading(doc, "3.1  方法概览", 2)
    add_para(doc,
        "RASP-LRM 是一种训练无关、轻量离线校准的双时间尺度 MLP 通道选择方法。离线时间尺度学习哪些通道在特定推理阶段中应被稳定保护；在线时间尺度根据当前阶段的近期中间激活，在保护核心之外选择实际可删除通道。方法不会将校准得到的完整 mask 冻结为部署策略，而是把校准结果转化为 protected prior。")
    add_callout(doc,
        "【图 3 待补：RASP-LRM 方法总览】",
        "建议采用左右双路径：左侧 Offline Calibration 显示独立混合校准池、正确且协议有效的四阶段轨迹、stage-conditioned WIFV-inspired prior / mean / protected core；右侧 Online Decoding 显示 causal stage parser、dense warmup、recent activation、non-core constrained selection、piecewise-constant mask cache、mean compensation 与 dense fallback。图中央突出：stage defines the safe search space, while the instance determines the actual mask。",
        fill=AMBER, border=AMBER_LINE)

    add_heading(doc, "3.2  显式阶段作为因果控制状态", 2)
    add_para(doc,
        "模型按固定顺序生成 SETUP、REASONING、VERIFY 与 FINAL 四个显式 marker。解码器只有在完整 marker 被观察后才更新 active stage，因此新状态从下一次 forward 开始生效，避免利用未来 token 或事后文本重建当前决策。若协议无法合法解析，后续生成回退到 dense MLP。与训练额外 stage classifier 相比，显式协议将控制信号直接嵌入模型输出，并提供可审计的失败处理路径。")

    add_heading(doc, "3.3  阶段条件保护先验，而非阶段固定策略", 2)
    add_para(doc,
        "在独立校准池中，我们仅保留答案正确、四阶段协议有效且未截断的结构化轨迹，并按阶段回放各层 MLP 中间激活。对每个 stage-layer pair，方法利用激活波动与 down-projection 列范数组合构造 WIFV-inspired importance，同时保存阶段均值用于输出补偿。高重要性通道形成 protected core，它们不参与在线删除候选。关键区别在于：被冻结的是阶段条件重要性、均值与保护边界，而不是一张完整的阶段 mask。")

    add_heading(doc, "3.4  保护核心之外的实例阶段动态选择", 2)
    add_para(doc,
        "在线解码时，每个 token 的 MLP intermediate activation 先进行 L2 归一化，并在当前阶段的最近窗口内聚合通道能量。对于 protected core 之外的通道，RASP-LRM 将离线 prior 与 recent activation 组合为可剪性分数：离线重要性较低且当前窗口活跃度较低的通道优先删除；若某个非核心通道在当前实例中持续活跃，recent evidence 会降低其删除优先级，从而实现 instance-level rescue。该过程使同一阶段中的实际 mask 可以随问题和局部推理状态变化，同时避免破坏稳定核心。")

    add_heading(doc, "3.5  稳定执行与安全边界", 2)
    add_para(doc,
        "为抑制短窗口噪声，Reasoning 与 Verify 在阶段开始后先执行 dense warmup；随后仅在 cache 缺失或达到阶段 refresh 间隔时更新 mask，使同一阶段内的决策呈分段常数。被 mask 的通道以阶段校准均值补偿其平均输出贡献，协议异常则触发 dense fallback。当前实现验证的是 decode-only logical MLP channel pruning：完整 gate/up/down 计算仍然执行，因此本文将主要结论限定为质量—逻辑稀疏权衡，而不把逻辑 mask 直接等同于已实现的物理加速。")

    add_heading(doc, "4  实验结果与证据闭环", 1)
    add_heading(doc, "4.1  近似匹配逻辑剪枝率下的主要结果", 2)
    add_para(doc,
        "我们在 Qwen3-1.7B 的显式阶段推理协议下比较 structured dense、trajectory-global static baseline 与 RASP-LRM。主要设置 t30_math_safe 在解码阶段实现约 34% reported logical MLP channel pruning。表 1 显示，在 GSM8K 与 MATH-500 共 1,819 道题上，RASP-LRM 的准确率为 63.77%，相比静态基线的 58.93% 提升 4.84 个百分点；两者 reported pruning 分别为 34.73% 与 35.27%。扩展到七个数据集 3,289 道题后，RASP-LRM 在 33.97% reported pruning 下达到 64.15%，相比静态基线在 34.08% 下的 60.75% 提升 3.41 个百分点，并多答对 112 道题。")
    add_result_table(doc)
    add_para(doc,
        "MATH-500 提供了尤其清晰的单数据集证据：RASP-LRM 的 reported pruning 为 36.21%，高于静态基线的 34.07%，但准确率仍由 42.60% 提升至 46.80%。因此，整体收益不能仅由动态方法剪得更少解释。结合表 2 的消融，现有证据更支持这样的解释：stage-conditioned protected prior 限制了高风险误删，而 current-instance activation 在剩余空间中修正了静态排序。")

    add_heading(doc, "4.2  结果如何回应 Motivation", 2)
    add_para(doc,
        "实验形成了从现象到方法的连续证据链。反事实诊断证明推理剪枝风险并非全程均匀；fixed stage-specific 消融表明阶段标签不足以唯一决定子网络；完整方法相对 fixed global 与 fixed stage 均取得更高准确率；大规模 full evaluation 则表明这一优势能够在近似匹配的 reported logical pruning 下保持。由此，RASP-LRM 的主要贡献不是增加一套阶段 mask bank，而是把阶段校准重新定位为受约束动态选择的安全先验。")

    add_callout(doc,
        "【图 4 待补：Accuracy–Reported Logical Pruning trade-off】",
        "建议将 dense、t20 与 t30 的 static / RASP-LRM 结果画成质量—逻辑稀疏曲线，分别标出核心数学集合和七数据集合计。横轴必须写为 Reported decode-only logical MLP channel pruning，而不是 Speedup；图注需注明当前仅包含 generation、未包含 dense prefill。",
        fill=AMBER, border=AMBER_LINE)

    add_heading(doc, "4.3  适用边界", 2)
    add_para(doc,
        "当前实验回答的是算法层面的质量—逻辑稀疏权衡：RASP-LRM 明显缩小了约 34% 通道 masking 相对静态剪枝造成的准确率损失，但尚未消除其相对 dense 模型的全部性能差距。由于当前实现未将动态 mask 转换为 reduced-width physical kernels，本文不报告端到端加速结论。该边界不影响本文关于阶段条件先验与实例动态选择的算法发现，但物理内核集成与系统测量属于后续部署研究。")

    add_heading(doc, "5  结论", 1)
    add_para(doc,
        "本文关注长链推理中被现有结构化剪枝忽略的轨迹内部非平稳性。反事实诊断表明，局部结构删除对最终答案的影响随推理功能、模块和强度变化；进一步消融则揭示了一个更关键的反常识事实：stage-aware 并不意味着为每个阶段冻结一张 mask。基于此，我们将 reasoning stage 重新定义为可验证的安全控制条件，并提出 RASP-LRM，以 stage-conditioned protected prior 提供稳定骨架、以 recent instance activation 决定非核心空间中的实际删除集合。七数据集结果表明，在约 34% reported decode-only logical MLP pruning 下，该方法相对近似匹配的静态基线取得稳定的准确率优势。更广泛地看，这项研究提供了一条不同于“全局静态”或“完全自由在线更新”的设计原则：动态模型压缩可以先用结构化先验限定低风险搜索空间，再在该空间内响应当前生成状态。")

    add_heading(doc, "参考文献（当前稿）", 1)
    refs = [
        "[1] An, Y., Zhao, X., Yu, T., Tang, M., and Wang, J. Fluctuation-Based Adaptive Structured Pruning for Large Language Models. AAAI, 2024.",
        "[2] Dong, H., Chen, B., and Chi, Y. Prompt-prompted Adaptive Structured Pruning for Efficient LLM Generation. COLM, 2024.",
        "[3] Le, Q., Diao, E., Wang, Z., et al. Probe Pruning: Accelerating LLMs through Dynamic Pruning via Model-Probing. ICLR, 2025.",
        "[4] GLASS: Global-Local Aggregation for Inference-time Sparsification of LLMs. arXiv:2508.14302, 2025.",
        "[5] Liang, X., Wang, H., Lai, H., et al. SEAP: Sparse Expert Activation Pruning Unlocks the Brainpower of Large Language Models. AAAI, 2026.",
        "[6] Ji, Y. and Sun, Y. OCP: Outlier-Centric Probing for Dynamic Structured Pruning of LLMs. ACL, 2026.",
    ]
    for ref in refs:
        p = add_para(doc, ref, first_indent=False, after=3, line=1.08)
        p.paragraph_format.left_indent = Pt(18)
        p.paragraph_format.first_line_indent = Pt(-18)
        for run in p.runs:
            set_run_font(run, size=8.7, color="303840")

    props = doc.core_properties
    props.title = "RASP-LRM Motivation, Insight and Method Story"
    props.subject = "AAAI Chinese manuscript draft"
    props.author = ""
    props.keywords = "RASP-LRM; structured pruning; reasoning stages; dynamic pruning"
    props.comments = "Generated as an evidence-aligned manuscript draft."

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
