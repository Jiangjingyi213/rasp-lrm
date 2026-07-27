#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph


SRC = Path(
    "/Users/jiangjiang/Library/Containers/com.tencent.xinWeChat/Data/Documents/"
    "xwechat_files/wxid_sh5wgza641lg12_d1c1/msg/file/2026-07/前述(1).docx"
)
FIG = Path(
    "/var/folders/4m/cg6tpxv51ss_gcjrtkzk2jbc0000gn/T/"
    "codex-clipboard-eb97aa5e-3c26-4080-8b31-aa0183e4f9cc.png"
)
OUT = Path("/Users/jiangjiang/Documents/Research/rasp-lrm/docs/前述_RASP示例图修订版.docx")


EXPLANATION = (
    "图 3 进一步以一个直观例子说明静态剪枝与 RASP-LRM 的区别。"
    "传统静态方法在完整推理轨迹中使用同一套全局 mask 和固定剪枝率，"
    "因此无法区分不同 reasoning stage 的通道需求：在某些阶段可能剪掉后续推理仍需要的关键通道，"
    "在另一些阶段又可能保留冗余通道或过度压缩当前状态，最终造成推理链条中间断裂。"
    "相比之下，RASP-LRM 不把阶段信息简单冻结为固定 mask，而是在不同阶段采用不同的剪枝强度与阶段化 mask，"
    "并通过 protected core 保护阶段关键通道；这样，runtime activation 仍可根据当前实例调整通道选择，"
    "同时阶段约束能够降低无约束动态剪枝误剪推理路径的风险。"
)

CAPTION = (
    "图 3 | 静态剪枝与 RASP-LRM 的阶段感知动态剪枝对比。"
    "静态剪枝在所有推理阶段使用固定 mask 和固定剪枝率，可能产生错剪、漏剪或多剪，"
    "从而破坏中间推理路径。RASP-LRM 根据当前 reasoning stage 调整剪枝率和阶段 mask，"
    "并使用 protected core 保留阶段关键通道，使动态剪枝在适应当前实例的同时维持推理连续性。"
)


def insert_paragraph_after(paragraph: Paragraph, text: str = "", style=None) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if style is not None:
        new_para.style = style
    if text:
        new_para.add_run(text)
    return new_para


def set_run_font(run, east_asia="SimSun", latin="Times New Roman", size=None):
    run.font.name = latin
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), latin)
    r_fonts.set(qn("w:hAnsi"), latin)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), latin)
    if size is not None:
        run.font.size = Pt(size)


def format_normal_para(p: Paragraph):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        set_run_font(run, size=10.5)


def format_caption(p: Paragraph):
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    for run in p.runs:
        set_run_font(run, size=9)


def main():
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if not FIG.exists():
        raise FileNotFoundError(FIG)

    doc = Document(str(SRC))

    # Insert after the insight paragraph beginning with “这些观察揭示”.
    target_idx = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("这些观察揭示"):
            target_idx = i
            break
    if target_idx is None:
        raise RuntimeError("Could not find insertion point after insight paragraph.")

    target = doc.paragraphs[target_idx]

    exp_para = insert_paragraph_after(target, EXPLANATION)
    exp_para.style = target.style
    format_normal_para(exp_para)

    img_para = insert_paragraph_after(exp_para)
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(4)
    img_para.paragraph_format.space_after = Pt(2)
    img_para.add_run().add_picture(str(FIG), width=Inches(4.95))

    caption_para = insert_paragraph_after(img_para, CAPTION)
    try:
        caption_para.style = doc.styles["Caption"]
    except KeyError:
        caption_para.style = target.style
    format_caption(caption_para)

    doc.save(str(OUT))
    print(OUT)


if __name__ == "__main__":
    main()
