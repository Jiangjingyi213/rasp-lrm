#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path("/Users/jiangjiang/Documents/Research/rasp-lrm")
OUT = ROOT / "docs" / "RASP_LRM_消融实验_正式论文插入版_主文精选.docx"
MD_OUT = ROOT / "docs" / "RASP_LRM_消融实验_正式论文插入版_主文精选.md"

CN_FONT = "SimSun"
EN_FONT = "Times New Roman"
BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(75, 75, 75)


SECTION_TITLE = "消融实验"

PARAGRAPHS = [
    "为了验证 RASP-LRM 各设计的作用，我们在五个推理数据集的 1040 个样本上进行核心消融。所有剪枝方法均采用相同的显式阶段化推理协议，并在约 33%--36% 的实际 MLP 通道剪枝率下比较。我们选取三类直接对应本文设计的对照：固定全局或固定阶段 mask、仅使用离线阶段先验的剪枝策略，以及去除 protected core 的动态策略。",
    "表 X 显示，完整 RASP-LRM 在核心消融设置中取得最高准确率。相比使用单一 trajectory-global mask 的 Static，RASP-LRM 在相近剪枝率下将准确率从 63.37% 提升到 65.58%，同时将截断率从 3.46% 降低到 1.83%。Fixed-Global 和 Fixed-Stage 分别只有 62.21% 和 61.54%，说明仅依赖离线固定 mask，甚至按阶段切换固定 mask，都不足以刻画长链推理过程中不断变化的 MLP 通道重要性。",
    "进一步地，Prior-Only 的准确率为 63.85%，但 fallback 率达到 84.23%，表明离线校准先验不能单独承担在线剪枝决策；当前样本的 runtime activation 对实例级通道选择是必要的。去除 protected core 后，准确率从 65.58% 降至 65.19%，截断率从 1.83% 升至 2.40%，说明 protected core 主要承担稳定动态 mask 的安全作用。总体来看，消融结果支持本文的核心设计：阶段先验提供安全边界，runtime activation 提供实例级修正，protected core 则降低长推理中的误剪风险。",
]

TABLE_CAPTION = (
    "表 X  核心消融结果。所有方法均在同一 1040 个样本上评测；"
    "Prune 表示实际 MLP 通道剪枝率，Δ 表示相对 Static 的准确率变化。"
)

TABLE_HEADERS = ["Variant", "Prior / mask", "Runtime", "Core", "Acc. ↑", "Prune", "Fallback", "Trunc.", "Δ"]
TABLE_ROWS = [
    ["Dense", "--", "--", "--", "76.73", "0.00", "6.25", "0.19", "--"],
    ["Static", "global fixed", "No", "No", "63.37", "34.06", "16.06", "3.46", "0.00"],
    ["Fixed-Global", "global fixed", "No", "No", "62.21", "36.10", "17.12", "5.19", "-1.16"],
    ["Fixed-Stage", "stage fixed", "No", "No", "61.54", "35.66", "18.46", "4.62", "-1.83"],
    ["Prior-Only", "stage prior", "No", "Yes", "63.85", "32.61", "84.23", "7.02", "+0.48"],
    ["No-Core", "stage prior", "Yes", "No", "65.19", "33.56", "18.65", "2.40", "+1.82"],
    ["RASP-LRM", "stage prior", "Yes", "Yes", "65.58", "33.44", "19.04", "1.83", "+2.21"],
]

NOTE = (
    "Note. Prior-Only 的高 fallback 率说明其准确率主要受到安全回退路径影响，"
    "因此它用于诊断离线先验的不足，而不是作为稳定在线剪枝方案。"
)


def set_run_font(run, en_font: str = EN_FONT, cn_font: str = CN_FONT) -> None:
    run.font.name = en_font
    r_fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), en_font)
    r_fonts.set(qn("w:hAnsi"), en_font)
    r_fonts.set(qn("w:eastAsia"), cn_font)
    r_fonts.set(qn("w:cs"), en_font)


def set_style_font(style, en_font: str = EN_FONT, cn_font: str = CN_FONT) -> None:
    style.font.name = en_font
    r_fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    r_fonts.set(qn("w:ascii"), en_font)
    r_fonts.set(qn("w:hAnsi"), en_font)
    r_fonts.set(qn("w:eastAsia"), cn_font)
    r_fonts.set(qn("w:cs"), en_font)


def add_run(paragraph, text: str, bold=False, italic=False, size: float | None = None, color=BLACK):
    r = paragraph.add_run(text)
    set_run_font(r)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = color
    if size is not None:
        r.font.size = Pt(size)
    return r


def add_para(doc, text: str = "", size: float = 10.3, after: float = 5, line: float = 1.14):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if text:
        add_run(p, text, size=size)
    return p


def set_cell_margins(cell, top=60, start=70, bottom=60, end=70):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        tag = f"w:{edge}"
        element = tc_borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            tc_borders.append(element)
        if edge_data is None:
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), edge_data.get("val", "single"))
            element.set(qn("w:sz"), edge_data.get("sz", "6"))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), edge_data.get("color", "000000"))


def set_cell_text(cell, text: str, bold=False, size=7.7, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    add_run(p, text, bold=bold, size=size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal = doc.styles["Normal"]
    set_style_font(normal)
    normal.font.size = Pt(10.3)
    normal.font.color.rgb = BLACK
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.14

    h1 = doc.styles["Heading 1"]
    set_style_font(h1)
    h1.font.size = Pt(14)
    h1.font.bold = True
    h1.font.color.rgb = BLACK
    h1.paragraph_format.space_before = Pt(6)
    h1.paragraph_format.space_after = Pt(5)


def add_booktabs_table(doc: Document) -> None:
    cap = add_para(doc, after=2, line=1.0)
    add_run(cap, TABLE_CAPTION, bold=False, size=8.8)

    table = doc.add_table(rows=1, cols=len(TABLE_HEADERS))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(table, [2.35, 2.15, 1.25, 1.05, 1.1, 1.1, 1.25, 1.05, 0.9])

    top_rule = {"val": "single", "sz": "12", "color": "000000"}
    mid_rule = {"val": "single", "sz": "8", "color": "000000"}
    bottom_rule = {"val": "single", "sz": "12", "color": "000000"}

    for j, h in enumerate(TABLE_HEADERS):
        cell = table.rows[0].cells[j]
        set_cell_text(cell, h, bold=True, size=7.6)
        set_cell_border(cell, top=top_rule, bottom=mid_rule)

    for i, row in enumerate(TABLE_ROWS, start=1):
        cells = table.add_row().cells
        is_last = i == len(TABLE_ROWS)
        is_main = row[0] == "RASP-LRM"
        for j, item in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j in (0, 1) else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(cells[j], item, bold=is_main, size=7.6, align=align)
            set_cell_border(cells[j], bottom=bottom_rule if is_last else None)

    note = add_para(doc, after=0, line=1.05)
    add_run(note, NOTE, italic=True, size=8.2, color=GRAY)


def build_markdown() -> str:
    lines = [f"## {SECTION_TITLE}", ""]
    for p in PARAGRAPHS:
        lines.extend([p, ""])
    lines.extend([TABLE_CAPTION, ""])
    lines.append("| " + " | ".join(TABLE_HEADERS) + " |")
    lines.append("|" + "|".join(["---"] * len(TABLE_HEADERS)) + "|")
    for row in TABLE_ROWS:
        lines.append("| " + " | ".join(row) + " |")
    lines.extend(["", NOTE, ""])
    return "\n".join(lines)


def main() -> None:
    doc = Document()
    configure_doc(doc)

    h = doc.add_paragraph(style="Heading 1")
    add_run(h, SECTION_TITLE, bold=True, size=14)

    for p in PARAGRAPHS:
        add_para(doc, p)

    add_booktabs_table(doc)

    doc.save(OUT)
    MD_OUT.write_text(build_markdown(), encoding="utf-8")
    print(OUT)
    print(MD_OUT)


if __name__ == "__main__":
    main()
