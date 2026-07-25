"""Build the clean Chinese Methodology revision with editable OMML equations.

This is a source artifact for the final DOCX.  Equations are created as Office
Math (OMML), not as images or plain-text substitutes.
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUT = Path(__file__).resolve().parents[1] / "docs" / "RASP_LRM_方法描述_清晰修订版.docx"
CN_FONT = "Arial Unicode MS"  # Stable in both Word and the LibreOffice QA renderer.


def set_run_font(run, name: str, size: float, *, bold: bool = False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor(*color)


def set_cell_width(cell, width_cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_cm * 567)))
    tc_w.set(qn("w:type"), "dxa")


def clear_cell_borders(cell):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "nil")
        borders.append(tag)
    tc_pr.append(borders)


def m(tag: str):
    return OxmlElement(f"m:{tag}")


def mr(text: str):
    run = m("r")
    t = m("t")
    t.text = text
    run.append(t)
    return run


def seq(*parts):
    out = []
    for part in parts:
        if isinstance(part, str):
            out.append(mr(part))
        elif isinstance(part, (list, tuple)):
            out.extend(seq(*part))
        else:
            out.append(part)
    return out


def sub(base, lower):
    node = m("sSub")
    e, s = m("e"), m("sub")
    e.extend(seq(base))
    s.extend(seq(lower))
    node.extend([e, s])
    return node


def sup(base, upper):
    node = m("sSup")
    e, s = m("e"), m("sup")
    e.extend(seq(base))
    s.extend(seq(upper))
    node.extend([e, s])
    return node


def subsup(base, lower, upper):
    node = m("sSubSup")
    e, lo, hi = m("e"), m("sub"), m("sup")
    e.extend(seq(base))
    lo.extend(seq(lower))
    hi.extend(seq(upper))
    node.extend([e, lo, hi])
    return node


def frac(num, den):
    node = m("f")
    n, d = m("num"), m("den")
    n.extend(seq(num))
    d.extend(seq(den))
    node.extend([n, d])
    return node


def rad(content):
    node = m("rad")
    pr = m("radPr")
    hide = m("degHide")
    hide.set(qn("m:val"), "1")
    pr.append(hide)
    deg, e = m("deg"), m("e")
    e.extend(seq(content))
    node.extend([pr, deg, e])
    return node


def topk(subscript, body):
    return sub("TopK", subscript), "(", body, ")"


def add_equation(doc: Document, parts, number: int | str):
    """Add a robust two-cell equation line; the central cell holds real OMML."""
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = (0.5, 15.2, 1.2)
    # python-docx writes equal table-grid columns by default.  Cell widths alone
    # are not enough for LibreOffice/Word to lay out a wide central equation cell.
    grid_cols = table._tbl.tblGrid.gridCol_lst
    for grid_col, width in zip(grid_cols, widths):
        grid_col.set(qn("w:w"), str(int(width * 567)))
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")
    for cell, width in zip(table.rows[0].cells, widths):
        set_cell_width(cell, width)
        clear_cell_borders(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    left, middle, right = table.rows[0].cells
    p = middle.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    omath = m("oMath")
    omath.extend(seq(*parts))
    p._p.append(omath)
    p_num = right.paragraphs[0]
    p_num.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_num.paragraph_format.space_after = Pt(1)
    r = p_num.add_run(f"({number})")
    set_run_font(r, "Times New Roman", 10.5)
    return table


def add_body(doc: Document, text: str, *, indent=True, after=5, before=0):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.15
    pf.space_before = Pt(before)
    pf.space_after = Pt(min(after, 3))
    if indent:
        pf.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    set_run_font(run, CN_FONT, 10.5)
    return p


def add_symbol_note(doc: Document, text: str):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
    pf.space_before = Pt(0)
    pf.space_after = Pt(3)
    pf.left_indent = Cm(0.74)
    run = p.add_run("式中，")
    set_run_font(run, CN_FONT, 9.5, bold=True, color=(64, 88, 102))
    run = p.add_run(text)
    set_run_font(run, CN_FONT, 9.5, color=(64, 88, 102))
    return p


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.keep_with_next = True
    if level == 1:
        pf.space_before = Pt(10)
        pf.space_after = Pt(6)
        size = 15
        color = (28, 43, 55)
    else:
        pf.space_before = Pt(8)
        pf.space_after = Pt(4)
        size = 12.5
        color = (28, 43, 55)
    r = p.add_run(text)
    set_run_font(r, CN_FONT, size, bold=True, color=color)
    return p


def add_title(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(1)
    r = p.add_run("METHODOLOGY")
    set_run_font(r, "Arial", 22, bold=True, color=(24, 35, 43))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("显式推理阶段驱动的安全动态结构化 MLP 通道剪枝")
    set_run_font(r, CN_FONT, 11, bold=True, color=(60, 84, 100))


def add_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("RASP-LRM  |  Methodology")
    set_run_font(r, "Arial", 8, color=(120, 130, 136))


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.4)
    section.bottom_margin = Cm(1.4)
    section.left_margin = Cm(1.9)
    section.right_margin = Cm(1.9)
    add_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = CN_FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), CN_FONT)
    normal.font.size = Pt(10.5)

    add_title(doc)
    add_heading(doc, "3  方法", 1)
    add_body(
        doc,
        "配对反事实剪枝分析表明：在输入、已生成前缀和剪枝强度相同的条件下，"
        "相同的结构扰动在不同推理阶段会带来不同的答案反转率。这意味着，推理过程中的"
        "计算冗余并不是固定不变的。全局静态掩码只能给出一套通道排序，无法反映这种变化；"
        "仅在 prompt 后作一次选择的策略，也无法利用后续推理过程中出现的实例证据。"
        "因此，RASP-LRM 以显式推理阶段为条件，在共享的稠密模型上为当前实例动态生成 MLP 通道掩码。"
        "它不为每个问题保存或训练一个新的剪枝模型：所有实例共享同一组模型参数和离线校准库，"
        "但各自维护独立的阶段状态、近期激活和掩码缓存。",
        after=8,
    )

    add_heading(doc, "3.1  问题设置与方法概览", 2)
    add_body(
        doc,
        "给定问题 q，语言模型以自回归方式生成 token 序列。记 l 为 Transformer 层编号，"
        "t 为当前解码步，x_{l,t} 为该层 MLP 的输入。对于 gated MLP，其进入 down projection "
        "之前的中间激活为："
    )
    add_equation(doc, [sub("h", "l,t"), " = ", "φ(", subsup("W", "l", "g"), " ", sub("x", "l,t"), ") ⊙ (", subsup("W", "l", "u"), " ", sub("x", "l,t"), ")"], 1)
    add_symbol_note(
        doc,
        "W_l^g、W_l^u 分别是 gate projection 和 up projection 的权重，φ(·) 为门控激活函数，"
        "⊙ 表示逐元素乘法，h_{l,t}∈R^{C_l} 为长度为 C_l 的中间通道向量。"
        "第 j 个结构化通道同时对应 gate/up projection 的第 j 个输出维度，以及 down projection 的第 j 个输入维度；"
        "因此，将该通道置零不会破坏张量维度。"
    )
    add_body(
        doc,
        "RASP 为每一层构造二值掩码 m_{s,l}(t)∈{0,1}^{C_l}，其中 s 表示当前推理阶段。"
        "m 的第 j 个元素为 0 时，通道 j 在当前 MLP 调用中不参与 down projection；为 1 时则保留。"
        "阶段 s 决定该层采用何种保护强度、剪枝预算和更新节奏；当前实例的近期激活决定在该约束下"
        "究竟删除哪些非关键通道。因此，掩码既可随阶段改变，也可在同一阶段因不同实例而不同。"
    )
    add_body(
        doc,
        "为使阶段在运行时可观察，我们在解码协议中使用 setup、reasoning、verify 和 final 四个阶段标记。"
        "解析器仅在观察到完整标记后更新当前阶段，因而不会依赖额外的阶段分类器。若标记序列不合法或"
        "无法确定活动阶段，系统立即回退为稠密 MLP。这个回退规则的作用是：当阶段信息不可信时，"
        "不再进行可能损害答案的动态剪枝。"
    )

    add_heading(doc, "3.2  阶段条件的离线校准", 2)
    add_body(
        doc,
        "离线校准的目标不是预先存储四张固定掩码，而是为每个阶段提供稳定的保护依据。"
        "我们在与正式测试集隔离的混合推理校准集上生成带阶段标记的轨迹，只保留答案正确、协议有效且"
        "未被最大长度截断的样本。统计时仅使用相邻阶段标记之间的内容 token；prompt 和阶段标记本身"
        "不进入阶段统计。对阶段 s、层 l，记收集到的有效 token 集为 T_{s,l}，则通道均值为："
    )
    add_equation(doc, [sub("μ", "s,l,j"), " = ", frac("1", ["|", sub("T", "s,l"), "|"]), sub("∑", ["t∈", sub("T", "s,l")]), sub("h", "l,t,j")], 2)
    add_symbol_note(
        doc,
        "h_{l,t,j} 是 h_{l,t} 的第 j 个分量，μ_{s,l,j} 是该通道在阶段 s 的经验均值，"
        "|T_{s,l}| 表示该阶段、该层参与统计的 token 数。该均值随后用于输出补偿，而不用于预测当前 token 的精确激活。"
    )
    add_body(
        doc,
        "我们以阶段内激活波动和该通道对层输出的影响共同定义离线重要性。具体地，"
        "将 down projection 中与通道 j 相连的权重列记为 W^d_{l,:,j}，定义："
    )
    add_equation(
        doc,
        [sub("V", "s,l,j"), " = ", frac("1", ["|", sub("T", "s,l"), "|"]), sub("∑", ["t∈", sub("T", "s,l")]), sup(["(", sub("h", "l,t,j"), " − ", sub("μ", "s,l,j"), ")"], "2")],
        "3a",
    )
    add_equation(
        doc,
        [sub("I", "s,l,j"), " = ", sub("V", "s,l,j"), " · ", subsup(["||", subsup("W", "l,:,j", "d"), "||"], "2", "2")],
        "3b",
    )
    add_symbol_note(
        doc,
        "V_{s,l,j} 是阶段内经验方差；||W^d_{l,:,j}||_2^2 是该权重列的平方 L2 范数；"
        "I_{s,l,j} 为阶段条件重要性。较大的 I 表示：该通道在该阶段的波动更明显，且其变化更容易"
        "传递到该层输出。I 只表示离线保护先验，不等同于“当前实例一定会使用该通道”。"
    )
    add_body(
        doc,
        "校准库对每个 stage-layer pair 仅保存重要性向量 I_{s,l} 和均值向量 μ_{s,l}。"
        "换言之，离线部分保存的是“该阶段哪些通道通常不应轻易删除”的知识；实际删除集合仍在推理时由当前实例决定。"
    )

    add_heading(doc, "3.3  基于近期激活的实例自适应选择", 2)
    add_body(
        doc,
        "仅依靠阶段先验仍然不够：即使两个问题处于同一阶段，它们也可能调用不同的通道。"
        "因此，RASP 在每次 MLP 前向中读取已经得到的中间激活 h，并在当前阶段的近期窗口内计算在线分数。"
        "为避免激活向量整体尺度主导评分，先对每个 token 的中间激活做 L2 归一化："
    )
    add_equation(doc, [sub("a", "l,t,j"), " = ", frac(sub("h", "l,t,j"), [sub(["||", sub("h", "l,t"), "||"], "2"), " + ε"])], 4)
    add_symbol_note(
        doc,
        "a_{l,t,j} 是归一化后通道 j 的激活值，ε 是避免分母为零的很小常数。"
        "归一化只用于排序：它比较一个 token 内各通道的相对活跃程度，而不改变原始 MLP 的计算。"
    )
    add_body(doc, "记 W_{s,l}(t) 为时刻 t 已经可见的当前阶段窗口，通道的近期激活分数为：")
    add_equation(doc, [sub("R", "s,l,j"), "(t) = ", rad([sub("∑", ["τ∈", sub("W", "s,l"), "(t)"]), sup(sub("a", "l,τ,j"), "2")])], 5)
    add_symbol_note(
        doc,
        "R_{s,l,j}(t) 汇总窗口内通道 j 的归一化能量。窗口仅包含当前步及其之前的激活，不访问未来 token；"
        "因而在线选择符合自回归解码的因果顺序。较大的 R 表示该通道在当前实例的最近推理上下文中持续活跃。"
    )
    add_body(
        doc,
        "在线信号可能因为短暂波动而不稳定。为避免把阶段中长期重要的结构误删，"
        "我们先从离线先验中选出硬保护集合："
    )
    add_equation(doc, [sub("Ω", "s,l"), " = ", *topk(["round(", sub("κ", "s"), sub("C", "l"), ")"], sub("I", "s,l"))], 6)
    add_symbol_note(
        doc,
        "κ_s 为阶段 s 的保护比例，TopK 返回分数最高的通道索引集合，Ω_{s,l} 中的通道永不参与动态删除。"
        "这一步把阶段稳定的重要通道作为安全下界保留下来，使近期激活只在剩余候选中调整排序。"
    )
    add_body(doc, "对于保护集合外的通道，我们在同一层内标准化离线与在线分数，并定义可剪性：")
    add_equation(
        doc,
        ["P(t) = ", sub("λ", "p"), " z(−I) + ", sub("λ", "r"), " z(−R(t))"],
        7,
    )
    add_symbol_note(
        doc,
        "此式为简洁起见省略 s、l、j 下标；z(·) 表示固定层 l 的所有通道上计算的 z-score，"
        "λ_p 和 λ_r 分别控制离线先验与在线证据的权重。"
        "负号使“重要性低、近期活跃度低”的通道得到较大的 P；P 越大，越优先删除。"
        "若一个非核心通道在当前实例中持续活跃，R 项会降低其 P，从而把它从删除候选中救回。"
    )
    add_body(doc, "设 r_s 为阶段 s 的名义剪枝率。为同时满足“至少保留一个通道”和“不得删除受保护通道”，实际可删除数写为：")
    add_equation(
        doc,
        ["k = min{", sub("C", "l"), " − 1, round(", sub("C", "l"), sub("r", "s"), ")}"],
        "8a",
    )
    add_equation(
        doc,
        [sup("k", "*"), " = min{k, ", sub("C", "l"), " − |", sub("Ω", "s,l"), "|}"],
        "8b",
    )
    add_symbol_note(
        doc,
        "此处省略 s、l 下标。k 是按名义预算得到的删除数，k^* 是考虑保护集合后的实际删除上限。"
        "round(·) 与实现中的通道计数规则一致。完整轨迹上的有效剪枝率还会受到稠密预填充、阶段 warmup 和安全回退影响，"
        "因此实验报告以实际测得的有效剪枝率为准。"
    )
    add_body(doc, "最终，RASP 从非保护通道中删除可剪性最高的 k̃_{s,l} 个通道：")
    add_equation(doc, ["D(t) = ", *topk(sup("k", "*"), ["{", sub("P", "j"), "(t) : j ∉ Ω}"])], 9)
    add_symbol_note(
        doc,
        "此式固定阶段 s 和层 l，D(t) 是实际删除的通道索引集合；m_{s,l,j}(t)=0 当且仅当 j∈D(t)，其余 m_{s,l,j}(t)=1。"
        "这说明 RASP 不是在四张预计算掩码之间切换：阶段规定选择的约束，当前实例的近期激活决定具体的删除集合。"
    )

    add_heading(doc, "3.4  因果执行与输出补偿", 2)
    add_body(
        doc,
        "推理开始时，prompt prefill 使用完整 MLP，并初始化阶段解析器、近期激活队列和掩码缓存。"
        "进入新阶段后，系统载入对应的先验、保护比例和剪枝预算。为避免阶段刚切换时只有极少上下文就作出剪枝决定，"
        "配置了 warmup 的阶段先以稠密 MLP 观察若干 token，同时更新近期窗口；之后仅在缓存缺失或到达刷新间隔时重新计算掩码。"
        "因此，掩码在一小段连续 token 内保持不变，而不是每步剧烈跳变。"
    )
    add_body(
        doc,
        "直接把被删除通道置零会同时删除其在校准数据上的平均输出贡献。我们用阶段条件均值作一阶补偿："
    )
    add_equation(
        doc,
        [sup("o", "mask"), " = ", subsup("W", "l", "d"), "[m ⊙ h + (1 − m) ⊙ ", sub("μ", "s,l"), "]"],
        10,
    )
    add_symbol_note(
        doc,
        "此式省略 l、t 下标。o^{mask} 是掩码后的 MLP 输出，m=m_{s,l}(t)，h=h_{l,t}，"
        "μ_{s,l} 是阶段 s、层 l 的通道均值向量。"
        "保留通道仍使用当前激活 h_{l,t}；被删除通道以其校准期平均激活替代。该补偿只减小系统性的均值偏移，"
        "不宣称能够重构当前 token 的真实激活。"
    )
    add_body(
        doc,
        "RASP 不训练额外路由器，也不更新基础模型参数。当前实现采用逻辑通道掩码：它验证的是通道选择策略本身，"
        "而不是端到端内核加速。在线解码按 batch size = 1 运行，每条推理轨迹独立维护状态，因而不同问题的激活不会相互混合。"
        "若未来扩展为批量解码，必须为每个样本分别维护阶段、近期窗口、warmup 计数器和掩码缓存。"
    )

    # A restrained closing boundary note; it prevents overclaiming speed-up.
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.left_indent = Cm(0.15)
    r = p.add_run("实现边界。")
    set_run_font(r, CN_FONT, 9.5, bold=True, color=(100, 66, 42))
    r = p.add_run(" 本文的“剪枝率”指逻辑 MLP 通道被屏蔽的比例；除非采用与掩码匹配的稀疏执行内核，"
                  "不将其直接解释为端到端推理时延的同比下降。")
    set_run_font(r, CN_FONT, 9.5, color=(100, 66, 42))

    # Add page numbering field.
    for section in doc.sections:
        p = section.footer.paragraphs[0]
        p.add_run("   ·   ")
        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        p._p.append(fld)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
